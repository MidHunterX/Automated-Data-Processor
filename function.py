from prompt_toolkit.completion import WordCompleter  # Autocompletion Engine
from sqlite3 import IntegrityError  # SQLite AccNo error
from prompt_toolkit import prompt   # Prompt for Autocompletion
from collections import Counter     # Most Common Value
from pandas import DataFrame        # Printing Tables
from pathlib import Path            # OS Independent filepath
from sys import exit
import threading    # Multithreading Stuff
import csv          # CSV file manipulation
import os           # Directory path support
import docx         # Docx parsing
import glob         # Finding files with extensions
import pdfplumber   # PDF parsing
import tabulate     # CLI Table Borders
import config as cfg
var = cfg.initVarCommon()


# =========================== [ @VAR_FUNCTIONS ] =========================== #


def loadIfscDataset(csv_file):
    """
    Parameter: CSV Dataset from RazorPay
    Returns: Dataset Dictionary loaded into memory

    dataset[row['IFSC']] = {
        'Bank': row['BANK'],
        'Branch': row['BRANCH'],
        'Centre': row['CENTRE'],
        'District': row['DISTRICT'],
        'State': row['STATE'],
        'Address': row['ADDRESS'],
        'City': row['CITY'],
    }
    """
    dataset = {}
    with open(csv_file, mode='r', encoding='utf-8') as file:
        reader = csv.DictReader(file)
        for row in reader:
            dataset[row['IFSC']] = {
                'Bank': row['BANK'],
                'Branch': row['BRANCH'],
                'Centre': row['CENTRE'],
                'District': row['DISTRICT'],
                'State': row['STATE'],
                'Address': row['ADDRESS'],
                'City': row['CITY'],
            }
    return dataset


def updateIfscInVar():
    var["ifsc_dataset"] = loadIfscDataset(var["ifsc_dataset_path"])


def getDistrictFromUser():

    csv_thread = threading.Thread(target=updateIfscInVar)
    csv_thread.start()

    # Get district while other thread works
    district = getDistrictInput()

    # Wait for updateIfscInVar() thread to finish
    csv_thread.join()

    return district


# ========================= [ @DATABASE_FUNCTIONS ] ========================= #


def writeToDB(conn, district, institution, student_data):
    """
    Arguments: (conn, district, institution, student_data)
        conn: Connection to database.db using sqlite3.connect()
        district: Name of District as String
        institution: Institution data from getInstitutionDetails function
        student_data: Student data from getStudentDetails function

    Returns:
        True: if inserted into DB successfully
        False: if any errors are encountered
    """

    try:
        cursor = conn.cursor()
        conn.execute("BEGIN TRANSACTION")

        # Insert Institution
        inst_name = institution["name"]
        inst_place = institution["place"]
        inst_number = institution["number"]
        inst_email = institution["email"]

        schoolSQL = """
        INSERT INTO Schools (
            SchoolName,
            District,
            Place,
            Phone,
            Email
        )
        VALUES ( ?, ?, ?, ?, ?)
        """
        values = inst_name, district, inst_place, inst_number, inst_email
        cursor.execute(schoolSQL, values)

        # Get the auto-incremented SchoolID
        school_id = cursor.lastrowid

        # Insert Students
        studentSQL = """
        INSERT INTO Students (
            SchoolID,
            StudentName,
            Class,
            IFSC,
            AccNo,
            AccHolder,
            Branch
        )
        VALUES ( ?, ?, ?, ?, ?, ?, ?)
        """

        for _, value in student_data.items():
            name = value[0]
            standard = value[1]
            ifsc = value[2]
            acc_no = value[3]
            holder = value[4]
            branch = value[5]
            variables = school_id, name, standard, ifsc, acc_no, holder, branch
            cursor.execute(studentSQL, variables)

        print("ℹ️ Commiting Changes")
        conn.commit()
        return True

    except IntegrityError as e:
        print(f"🔴 IntegrityError: {e}")
        conn.rollback()
        return False

    except Exception as e:
        print(f"Error: {e}")
        conn.rollback()
        return False


def getIndianState():
    indian_states = [
        "Andhra Pradesh",
        "Arunachal Pradesh",
        "Assam",
        "Bihar",
        "Chhattisgarh",
        "Goa",
        "Gujarat",
        "Haryana",
        "Himachal Pradesh",
        "Jharkhand",
        "Jammu and Kashmir",
        "Karnataka",
        "Kerala",
        "Madhya Pradesh",
        "Maharashtra",
        "Manipur",
        "Meghalaya",
        "Mizoram",
        "Nagaland",
        "Odisha",
        "Punjab",
        "Rajasthan",
        "Sikkim",
        "Tamil Nadu",
        "Telangana",
        "Tripura",
        "Uttar Pradesh",
        "Uttarakhand",
        "West Bengal",
        "Puducherry"
    ]

    state_completer = WordCompleter(indian_states)
    # Prompt for state selection
    selected_state = prompt('Enter a state: ', completer=state_completer)
    return selected_state


# ========================== [ @PRINT_FUNCTIONS ] ========================== #


def printInstitution(institution):
    inst_name = institution["name"]
    inst_place = institution["place"]
    inst_number = institution["number"]
    inst_email = institution["email"]
    print(f"{inst_name}\n{inst_place}\n{inst_number}\n{inst_email}")


def printStudentData(student_data):
    i = 1
    for _, value in student_data.items():
        name = value[0]
        standard = value[1]
        ifsc = value[2]
        acc_no = value[3]
        holder = value[4]
        branch = value[5]
        print(f"{i}: {name},{standard},{ifsc},{acc_no},{holder},{branch}")
        i += 1


def printStudentDataFrame(student_data):
    ifsc_dataset = var["ifsc_dataset"]
    student_record = []
    for key, val in student_data.items():

        ifsc = val[2]
        branch = getBranchFromIfsc(ifsc, ifsc_dataset)

        # Check IFSC Validity
        if branch == "":
            ifsc = f"{ifsc}❌"
        else:
            ifsc = f"{ifsc}✅"

        row = [key+1, val[0], val[1], ifsc, val[3], val[4], val[5]]
        student_record.append(row)

    df = DataFrame(
        student_record,
        columns=['', 'Name', 'Std', 'IFSC', 'Account No', 'Holder', 'Branch']
    )
    print(df.to_string(index=False))


# ======================= [ @VERIFICATION_FUNCTION ] ======================= #


def userVerifyStudentData(student_data):
    """
    Asks user for confirmation if student standard is valid
    else returns n
    """
    valid_std = isValidStudentStd(student_data)
    if valid_std is True:
        verification = input("Correct? (ret / n): ")
        if verification == "":
            return True
        else:
            return False
    else:
        # Moves to investigation_dir for checking
        print("⚠️ Cannot convert std to num equivalents :(")
        return False


# ======================== [ @PROCESSOR_FUNCTIONS ] ======================== #


def convertStdToNum(data):
    """
    Parameter: Student Standard / Class Number
    Returns: Numeric Value if String
    """
    data = str(data)
    data = data.strip().lower()

    std_dataset = {
        1: [
            "1",
            "1a",
            "1b",
            "1c",
            "1d",
            "1e",
            "1 a",
            "1 b",
            "1 c",
            "1 d",
            "1 e",
            "i",
            "1st",
            "one",
            "first"
        ],
        2: [
            "2",
            "2a",
            "2b",
            "2c",
            "2d",
            "2e",
            "2 a",
            "2 b",
            "2 c",
            "2 d",
            "2 e",
            "ii",
            "2nd",
            "two",
            "second"
        ],
        3: [
            "3",
            "3a",
            "3b",
            "3c",
            "3d",
            "3e",
            "3 a",
            "3 b",
            "3 c",
            "3 d",
            "3 e",
            "iii",
            "3rd",
            "three",
            "third"
        ],
        4: [
            "4",
            "4a",
            "4b",
            "4c",
            "4d",
            "4e",
            "4 a",
            "4 b",
            "4 c",
            "4 d",
            "4 e",
            "iv",
            "1v",
            "4th",
            "four",
            "fourth"
        ],
        5: [
            "5",
            "5a",
            "5b",
            "5c",
            "5d",
            "5e",
            "5 a",
            "5 b",
            "5 c",
            "5 d",
            "5 e",
            "v",
            "5th",
            "five",
            "fifth",
        ],
        6: [
            "6",
            "6a",
            "6b",
            "6c",
            "6d",
            "6e",
            "6 a",
            "6 b",
            "6 c",
            "6 d",
            "6 e",
            "vi",
            "v1",
            "six",
            "6th",
            "sixth",
        ],
        7: [
            "7",
            "7a",
            "7b",
            "7c",
            "7d",
            "7e",
            "7 a",
            "7 b",
            "7 c",
            "7 d",
            "7 e",
            "vii",
            "v11",
            "7th",
            "seven",
            "seventh",
        ],
        8: [
            "8",
            "8a",
            "8b",
            "8c",
            "8d",
            "8e",
            "8 a",
            "8 b",
            "8 c",
            "8 d",
            "8 e",
            "v111",
            "viii",
            "8th",
            "eight",
            "eighth",
        ],
        9: [
            "9",
            "9a",
            "9b",
            "9c",
            "9d",
            "9e",
            "9 a",
            "9 b",
            "9 c",
            "9 d",
            "9 e",
            "1x",
            "ix",
            "9th",
            "nine",
            "nineth",
        ],
        10: [
            "10",
            "10a",
            "10b",
            "10c",
            "10d",
            "10e",
            "10 a",
            "10 b",
            "10 c",
            "10 d",
            "10 e",
            "x",
            "10th",
            "ten",
            "tenth",
        ],
        11: [
            "11",
            "x1",
            "xi",
            "11th",
            "plus one",
            "plusone",
            "+1",
            "+1 science",
            "+1 commerce",
            "+1 humanities",
        ],
        12: [
            "12",
            "x11",
            "xii",
            "12th",
            "plus two",
            "plustwo",
            "+2",
            "+2 science",
            "+2 commerce",
            "+2 humanities",
        ],
        13: [
            "1 dc",
            "1dc",
            "i dc",
            "idc",
            "ist dc",
            "1stdc",
            "1st dc"
        ],
        14: [
            "2 dc",
            "2dc",
            "ii dc",
            "iidc",
            "iind dc",
            "2nddc",
            "2nd dc"
        ],
        15: [
            "3 dc",
            "3dc",
            "iii dc",
            "iiidc",
            "iiird dc",
            "3rddc",
            "3rd dc"
        ],
        16: [
            "1 pg",
            "1pg",
            "i pg",
            "ipg",
            "ist pg",
            "1st pg",
            "1stpg"
        ],
        17: [
            "2 pg",
            "2pg",
            "ii pg",
            "iipg",
            "iind pg",
            "2ndpg",
            "2nd pg"
        ],
    }
    if isinstance(data, str):
        data = data.lower()
        for key, values in std_dataset.items():
            for value in values:
                if data == value:
                    data = key
    return data


def getBranchFromIfsc(ifsc, ifsc_dataset):
    """
    Arguments: (ifsc, ifsc_dataset)
        - ifsc_list: IFSC Code
        - ifsc_dataset: IFSC Razorpay Dataset from loadIfscDataset()

    Returns:
        - "": If there exists no record of IFSC in dataset
        - branch: If Branch for IFSC is found
    """
    branch = ""
    if type(ifsc) is str:
        ifsc_details = ifsc_dataset.get(ifsc)
        if ifsc_details:
            branch = ifsc_details["Branch"]

    # Clean Up Branch
    if "IMPS" in branch:
        branch = branch.replace("IMPS", "").strip()

    return branch


def normalizeStudentData(student_data):
    """
    Parameter:
        - student_data: Student data from getStudentDetails()

    Returns: A dictionary of tuples with corrected student data

    data = {
        0: (name, standard, ifsc, acc_no, holder, branch),
        1: (name, standard, ifsc, acc_no, holder, branch),
        2: (name, standard, ifsc, acc_no, holder, branch)
    }
    """
    ifsc_dataset = var["ifsc_dataset"]
    i = 0
    data = {}
    for _, value in student_data.items():
        name = value[0]
        standard = value[1]
        ifsc = value[2]
        acc_no = value[3]
        holder = value[4]
        branch = value[5]

        # Empty Acc Holder fix
        if holder == "":
            holder = name

        # Normalizing Standard standard to Int variant
        standard = convertStdToNum(standard)

        # Normalizing Branch from IFSC using RazorPay Dataset
        rbi_branch = getBranchFromIfsc(ifsc, ifsc_dataset)

        # "," fix and long branch name fix
        if rbi_branch and "," not in rbi_branch and len(rbi_branch) < 30:
            branch = rbi_branch

        # Extracted data
        data[i] = name, standard, ifsc, acc_no, holder, branch
        i = i + 1

    return data


def get_most_common_value(a_list):
    """
    Parameters: A List of Values
    Returns: The most common value from list
    """
    count = Counter(a_list)
    mostCommon = count.most_common(1)
    return mostCommon[0][0]


def guessDistrictFromIfscList(ifsc_list):
    """
    Parameters: (ifsc_list)
    Returns: Guessed District as a String
    """
    district_list = []
    ifsc_dataset = var["ifsc_dataset"]
    # Create a list of Districts
    for ifsc in ifsc_list:
        district = getDistrictFromIfsc(ifsc, ifsc_dataset)
        district_list.append(district)

    # Finding the most occured District
    district = get_most_common_value(district_list)

    # If District is Unknown, log data for user to verify
    if district == "Unknown":
        print("Couldn't decide District :(")
        print(f"Guess Data: {district_list}\n")

    # If District is unrecognized district, select from most common 3 values
    district_dataset = var["district_dataset"]
    if district not in district_dataset:
        count = Counter(district_list)
        mostCommon_list = count.most_common()
        for i in mostCommon_list:
            if i[0] in district_dataset:
                district = i[0]
                break

    return district


def getDistrictFromIfsc(ifsc, ifsc_dataset):
    """
    Parameters: (ifsc_code, ifsc_dataset)
    Returns: District as a String
    """
    district = "Unknown"
    ifsc_info = ifsc_dataset.get(ifsc)
    district_list = var["district_dataset"]

    if ifsc_info:

        # District Finder (Initial Algorithm)
        district = ifsc_info["District"]
        # Return correct item from district data
        for item in district_list:
            if item.lower() == district.lower():
                district = item

        # District Finder v2.0 (New Algorithm)
        if district not in district_list:
            value_list = []
            for value in ifsc_info.values():
                value_list.append(value)
            district = get_most_common_value(value_list)

            # Return correct item from district data
            for item in district_list:
                if item.lower() == district.lower():
                    district = item

        # District Finder v1.0 (Fallback)
        if district not in district_list:
            for item in district_list:
                address = ifsc_info["Address"]
                if item.lower() in address.lower():
                    district = item

            # Return correct item from district data
            for item in district_list:
                if item.lower() == district.lower():
                    district = item

    return district


def getStudentIfscList(student_data):
    """
    Parameter: Student Data from getStudentDetails()
    Returns: Iterable List of IFSC code

    ifsc = ["ifsc1", "ifsc2", "ifsc3", "ifsc4"]
    """
    i = 0
    ifsc = []
    for _, value in student_data.items():
        ifsc.append(value[2])
        i += 1
    return ifsc


def cleanStudentData(student_data):
    """
    Parameter: Student Data from getStudentDetails()
    Returns: A dictionary of tuples with cleaned up values for processing

    data = {
        0: (name, standard, ifsc, acc_no, holder, branch),
        1: (name, standard, ifsc, acc_no, holder, branch),
        2: (name, standard, ifsc, acc_no, holder, branch)
    }
    """
    i = 0
    data = {}
    for _, value in student_data.items():
        name = value[0]
        standard = value[1]
        ifsc = value[2]
        acc_no = value[3]
        holder = value[4]
        branch = value[5]

        # Cleaning spaces and newline from data
        name = convertParagraphToLine(name)
        standard = convertParagraphToLine(standard)
        holder = convertParagraphToLine(holder)
        branch = convertParagraphToLine(branch)

        # Cleaning up important data
        acc_no = str(acc_no).strip()
        ifsc = str(ifsc).strip()
        name = name.strip()
        holder = holder.strip()

        # Extracted data
        data[i] = name, standard, ifsc, acc_no, holder, branch
        i = i + 1

    return data


def renameFilenameToInstitution(file, institution):
    file_path = Path(file)

    dir = file_path.parent
    extension = file_path.suffix

    new_name = institution["name"].replace(".", "").replace(",", "")
    new_name = f"{new_name}{extension}"

    new_path = dir / new_name
    file_path.rename(new_path)

    return new_path


# ========================== [ @DIFF_FUNCTIONS ] ========================== #


def checkExistingAccounts(data, cursor):
    """
    Parameters:
    - data (dict): output of getStudentDetails()
    - cursor (sqlite3.Cursor): SQLite database cursor object to execute queries.

    Returns:
    - bool: True if any account number is found, False otherwise.
    """

    acc_nos = [entry[3] for entry in data.values()]  # entry[3] is acc_no
    query = """
    SELECT COUNT(1)
    FROM Students
    WHERE AccNo IN ({})
    """.format(','.join('?' for _ in acc_nos))
    cursor.execute(query, acc_nos)
    result = cursor.fetchone()[0]
    return result > 0


def getExistingAccounts(studentDetails, cursor):
    """
    Parameters:
    - studentDetails (dict): output of getStudentDetails()
    - cursor (sqlite3.Cursor): SQLite database cursor object to execute queries.

    Returns:
    - list: A list of tuples containing existing accounts.
        data = [(db_class, db_name, db_acc, db_ifsc, db_branch)]
    """

    acc_nos = [entry[3] for entry in studentDetails.values()]  # entry[3] is acc_no
    query = """
    SELECT Class, StudentName, AccNo, IFSC, Branch
    FROM Students
    WHERE AccNo IN ({})
    """.format(','.join('?' for _ in acc_nos))

    cursor.execute(query, acc_nos)
    existing_accounts = cursor.fetchall()

    studentDetails = []
    for db_class, db_name, db_acc, db_ifsc, db_branch in existing_accounts:
        studentDetails.append(
            (db_class, db_name, db_acc, db_ifsc, db_branch)
        )

    return studentDetails


def getNonExistingAccounts(studentDetails, existingAccounts):
    """
    Parameters:
    - studentDetails (dict): output of getStudentDetails()
    - existingAccounts (list): output of getExistingAccounts()

    Returns:
    - list: A list of tuples containing newly recieved students.
        data = [(name, standard, ifsc, acc_no, holder, branch)]
    """
    existing_acc_nos = {entry[2] for entry in existingAccounts}

    non_existing_accounts = []

    for student in studentDetails.values():
        name = student[0]
        standard = student[1]
        ifsc = student[2]
        acc_no = student[3]
        holder = student[4]
        branch = student[5]

        if acc_no not in existing_acc_nos:
            non_existing_accounts.append((name, standard, ifsc, acc_no, holder, branch))

    return non_existing_accounts


def printExistingAccounts(comparison_list):
    """
    Parameters:
    - comparison_list (list): Output of getExistingAccounts()

    Prints:
    - Table showing database data and the provided data side by side.
    """

    df = DataFrame(comparison_list, columns=[
        'STD', 'Name', 'Acc No', 'IFSC', 'Branch',
    ])

    print(tabulate.tabulate(df, headers='keys', tablefmt='rounded_outline', showindex=False))


def printExistingAccountsDiff(studentData, existingAccounts):
    """
    Parameters:
    - existingAccounts (list): Output of getExistingAccounts()
    - studentData (dict): Output of getStudentDetails()

    Prints:
    - Table showing difference between database data and the provided data side by side.
    """

    ifsc_dataset = var["ifsc_dataset"]

    comparison_list = []
    data_by_acc = {entry[3]: entry for entry in studentData.values()}

    for db_class, db_name, db_acc, db_ifsc, db_branch in existingAccounts:
        if db_acc in data_by_acc:
            name = data_by_acc[db_acc][0]
            std = data_by_acc[db_acc][1]
            ifsc = data_by_acc[db_acc][2]
            accno = data_by_acc[db_acc][3]

            branch = getBranchFromIfsc(ifsc, ifsc_dataset)

            # Check IFSC Validity
            if branch == "":
                ifsc_check = f"{ifsc} ❌"
            else:
                ifsc_check = f"{ifsc} ✅"

            data_tuple = (
                f"{db_name} -> {name}" if db_name != name else name,
                f"{db_class} -> {std}" if db_class != std else std,
                f"{db_acc} -> {accno}" if db_acc != accno else accno,
                f"{db_ifsc} -> {ifsc_check}" if db_ifsc != ifsc else ifsc,
                f"{db_branch}"
            )

            # Append both db and data tuples for comparison
            comparison_list.append(data_tuple)

    df = DataFrame(comparison_list, columns=[
        'Name', 'STD', 'Acc No', 'IFSC', 'Branch',
    ])
    print(tabulate.tabulate(df, headers='keys', tablefmt='rounded_outline', showindex=False))

    newly_added_list = []
    nonExistingAccounts = getNonExistingAccounts(studentData, existingAccounts)
    for name, standard, ifsc, acc_no, holder, branch in nonExistingAccounts:
        branch = getBranchFromIfsc(ifsc, ifsc_dataset)
        if branch == "":
            ifsc = f"{ifsc} ❌"
        else:
            ifsc = f"{ifsc} ✅"
        newly_added_list.append((name, standard, ifsc, acc_no, holder, branch))

    if newly_added_list:
        new_df = DataFrame(newly_added_list, columns=[
            'Name', 'STD', 'IFSC', 'Acc No', 'Holder', 'Branch',
        ])
        print("ℹ️ New students in form:")
        print(tabulate.tabulate(new_df, headers='keys', tablefmt='rounded_outline', showindex=False))


# ========================== [ @PARSER_FUNCTIONS ] ========================== #


def getInstitutionDetails(file):
    """
    Parameters: Supported File
    Returns: Dictionary of Institution Details

    data = {
        "name": name_of_institution,
        "place": place,
        "number": phone_number,
        "email": email_id
    }
    """
    _, file_extension = os.path.basename(file).split(".")
    data = {}

    if file_extension == "docx":
        docx_file = file
        data = getInstitutionDetailsDocx(docx_file)

    if file_extension == "pdf":
        pdf_file = file
        data = getInstitutionDetailsPdf(pdf_file)

    return data


def getStudentDetails(file):
    """
    Parameter: Supported File
    Returns: A dictionary of tuples with Student details

    data = {
        0: (name, standard, ifsc, acc_no, holder, branch),
        1: (name, standard, ifsc, acc_no, holder, branch),
        2: (name, standard, ifsc, acc_no, holder, branch)
    }
    """
    _, file_extension = os.path.basename(file).split(".")
    data = {}

    if file_extension == "docx":
        docx_file = file
        data = getStudentDetailsDocx(docx_file)

    if file_extension == "pdf":
        pdf_file = file
        data = getStudentDetailsPdf(pdf_file)

    return data


def getInstitutionDetailsDocx(docx_file):
    """
    Parameters: Document.docx file
    Returns: Dictionary of Institution Details

    data = {
        "name": name_of_institution,
        "place": place,
        "number": phone_number,
        "email": email_id
    }
    """

    doc = docx.Document(docx_file)
    inside_institution_details = False
    name_of_institution = ""
    place = ""
    phone_number = ""
    email_id = ""

    for paragraph in doc.paragraphs:
        text = paragraph.text
        # Check if paragraph contains the "Institution Details"
        if text.startswith("Institution Details"):
            inside_institution_details = True
        elif inside_institution_details:
            # Split the paragraph by the colon
            parts = text.split(':')
            if len(parts) == 2:
                key = parts[0].strip()
                value = parts[1].strip()

                # Assign values to variables
                if key == "Name of the Institution":
                    name_of_institution = value
                elif key == "Place":
                    place = value
                elif key == "Phone number":
                    phone_number = value
                elif key == "Email Id":
                    email_id = value

    # Extracted data
    data = {
        "name": name_of_institution,
        "place": place,
        "number": phone_number,
        "email": email_id
    }
    return data


def getStudentDetailsDocx(docx_file):
    """
    Parameter: Document.docx file
    Returns: A dictionary of tuples with Student details

    data = {
        0: (name, standard, ifsc, acc_no, holder, branch),
        1: (name, standard, ifsc, acc_no, holder, branch),
        2: (name, standard, ifsc, acc_no, holder, branch)
    }
    """

    doc = docx.Document(docx_file)
    data = {}
    i = 0
    # Iterate through the tables in the document
    for table in doc.tables:
        for row in table.rows:
            first_column = row.cells[0].text
            if first_column != "" and first_column != "STUDENT NAME":
                name = row.cells[0].text
                standard = row.cells[1].text
                ifsc = row.cells[2].text
                acc_no = row.cells[3].text
                holder = row.cells[4].text
                branch = row.cells[5].text

                # Extracted data
                data[i] = name, standard, ifsc, acc_no, holder, branch
                i = i + 1
    return data


def getStudentDetailsPdf(pdf_file):
    """
    Parameter: PDF File
    Returns: A dictionary of tuples with Student details

    data = {
        0: (name, standard, ifsc, acc_no, holder, branch),
        1: (name, standard, ifsc, acc_no, holder, branch),
        2: (name, standard, ifsc, acc_no, holder, branch)
    }
    """
    with pdfplumber.open(pdf_file) as pdf:
        data = {}
        i = -1
        for page in pdf.pages:
            # Generate CSV list from PDF table
            table = page.extract_table()
            if table:
                for row in table:
                    # Replace \n substring with space
                    cleaned_row = []
                    for cell in row:
                        if isinstance(cell, str):
                            cleaned_row.append(cell.replace('\n', ' '))
                        else:
                            cleaned_row.append(cell)

                    name = cleaned_row[0]
                    standard = cleaned_row[1]
                    ifsc = cleaned_row[2]
                    acc_no = cleaned_row[3]
                    holder = cleaned_row[4]
                    branch = cleaned_row[5]

                    # Extracted data
                    if name:  # For avoiding empty rows
                        data[i] = name, standard, ifsc, acc_no, holder, branch
                        i = i + 1

    # Removes unwanted Header data
    data.pop(-1)

    return data


def getInstitutionDetailsPdf(pdf_file):
    """
    Parameters: Document.pdf file
    Returns: Dictionary of Institution Details

    data = {
        "name": name_of_institution,
        "place": place,
        "number": phone_number,
        "email": email_id
    }
    """
    institution_details = ""
    name_of_institution = ""
    place = ""
    phone_number = ""
    email_id = ""

    with pdfplumber.open(pdf_file) as pdf:
        for page in pdf.pages:
            # Extract PDF as text
            text = page.extract_text()
            if "Institution Details" in text:
                start = text.index("Name of the Institution")
                end = text.index("Student Details")
                institution_details = text[start:end]

    # Splitting text at '\n' into a list
    lines = institution_details.split('\n')

    for line in lines:
        parts = line.split(':')
        if len(parts) == 2:
            key = parts[0].strip()
            value = parts[1].strip()

            # Assign values to variables
            if key == "Name of the Institution":
                name_of_institution = value
            elif key == "Place":
                place = value
            elif key == "Phone number":
                phone_number = value
            elif key == "Email Id":
                email_id = value

    # Extracted data
    data = {
        "name": name_of_institution,
        "place": place,
        "number": phone_number,
        "email": email_id
    }
    return data


# ======================== [ @VALIDATION_FUNCTIONS ] ======================== #


def isValidStudentStd(student_data):
    """
    Parameter: Student Data from getStudentDetails()
    Returns: Boolean Value (True / False) if all standards are converted to int
    """
    data = []
    for _, value in student_data.items():
        standard = value[1]
        if type(standard) is int:
            data.append(True)
        else:
            data.append(False)

    return all(data)


def correctFormat(file):
    """
    Returns True if file is in correct Format
    """
    data = False
    try:
        _, file_extension = os.path.basename(file).split(".")

        if file_extension == "docx":
            docx_file = file
            data = correctDocxFormat(docx_file)

        if file_extension == "pdf":
            pdf_file = file
            data = correctPdfFormat(pdf_file)

    except ValueError:
        print("⚠️ ValueError: Possible dot in file name")

    return data


def correctPdfFormat(pdf_file):
    """
    Returns True if PDF is in correct Format
    """
    flags = {
        "Institution Heading": False,
        "Institution Lines": False,
        "Student Heading": False,
        "Student Table": False
    }
    try:
        with pdfplumber.open(pdf_file) as pdf:
            for page in pdf.pages:

                # ====== TEXT PARAGRAPH STARTS ====== #
                text = page.extract_text()

                # Check Heading: Institution Details
                if "Institution Details" in text:
                    flags["Institution Heading"] = True

                    # Check Length: Institution Details
                    start = text.index("Name of the Institution")
                    end = text.index("Student Details")
                    institution_details = text[start:end].splitlines()
                    if len(institution_details) == 4:
                        flags["Institution Lines"] = True

                # Check Heading: Student Details
                if "Student Details" in text:
                    flags["Student Heading"] = True

                # =========== TABLE STARTS =========== #
                table = page.extract_table()

                # Check Content: Student Table
                if table:
                    flags["Student Table"] = True
    except ValueError:
        pass

    # Logs
    if flags["Institution Heading"] is False:
        print("Heading not found: Institution Details")
    if flags["Institution Lines"] is False:
        print("Data Incomplete: Institution Details")
    if flags["Student Heading"] is False:
        print("Heading not found: Student Details")
    if flags["Student Table"] is False:
        print("Object not found: Student Table")

    status = all(flags.values())
    return status


def correctDocxFormat(docx_file):
    """
    Returns True if DOCX is in correct Format
    """
    inside_institution_details = False
    flags = {
        "name": False,
        "place": False,
        "number": False,
        "email": False,
    }
    doc = docx.Document(docx_file)
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if text.startswith("Institution Details"):
            inside_institution_details = True
        elif inside_institution_details:
            if text.startswith("Name of the Institution"):
                flags["name"] = True
            if text.startswith("Place"):
                flags["place"] = True
            if text.startswith("Phone number"):
                flags["number"] = True
            if text.startswith("Email Id"):
                flags["email"] = True

    # Logs
    if flags["name"] is False:
        print("Heading not found: Name of the Institution")
    if flags["place"] is False:
        print("Entry not found: Place")
    if flags["number"] is False:
        print("Entry not found: Number")
    if flags["email"] is False:
        print("Entry not found: Email")

    status = all(flags.values())
    return status


# ========================== [ @OTHER_FUNCTIONS ] ========================== #


def sanitizeFilename(file):
    """
    Parameter:
        - file (str): filename.extension
    Returns:
        - file (str): sanitized_filename.extension
    """
    try:
        file_name, file_extension = os.path.basename(file).split(".")
    except ValueError:
        path = Path(file)
        file_name = path.stem
        file_extension = path.suffix

        sanitized_file_name = file_name.replace(".", "")
        new_file_name = sanitized_file_name + file_extension
        new_file_path = path.with_name(new_file_name)
        path.rename(new_file_path)

        file = new_file_path
        print(f"📝 File renamed to: {file}")

    return file


def printTextBox_Centered(text):
    box_width = 80

    if len(text) > box_width - 4:
        text = text[:box_width - 7] + '...'

    freespace = box_width - len(text) - 4
    space = ' ' * (freespace // 2)

    horizontal_line = '+' + '-' * (box_width - 2) + '+'

    print(horizontal_line)
    print(f"|{space} {text}{space + ' ' if freespace % 2 != 0 else space}|")
    # print(f"|{' ' * (box_width - 2)}|")
    print(horizontal_line)


def printFileNameHeader(file):
    try:
        file_name, file_extension = os.path.basename(file).split(".")
        printTextBox_Centered(f"📄 {file_name}.{file_extension}")
    except ValueError:
        printTextBox_Centered(f"📄 {file}")


def getFileList(dir, extensions):
    """
    Parameters: (dir, extensions)
        - dir: Directory Path
        - extensions: List of File extensions
    Returns: A list of file path.

    file_list = [file1.ext1, file2.ext1, file3.ext2, file4.ext2]
    """

    file_list = []

    if not isinstance(extensions, list):
        ext = [extensions]

    for ext in extensions:
        supported_files = glob.glob(os.path.join(dir, f"*{ext}"))
        for file in supported_files:
            file_list = file_list + [file]

    return file_list


def initNestedDir(input_dir, nest_name):
    directory_path = os.path.join(input_dir, nest_name)
    if not os.path.exists(directory_path):
        os.mkdir(directory_path)
    return directory_path


def getDistrictInput():
    district = "Unknown"
    try:
        print("""
        1: TVM    6: IDK   11: KKD
        2: KLM    7: EKM   12: WYD
        3: PTA    8: TSR   13: KNR
        4: ALP    9: PKD   14: KSD
        5: KTM   10: MLP    0: Unknown
        """)
        district_dataset = var["district_dataset"]
        data = int(input("Enter District No: "))
        data -= 1
        if data <= 13 and data >= 0:
            district = district_dataset[data]
        return district

    except KeyboardInterrupt:
        print("Caught the Interrupt ;)")
        exit()

    except ValueError:
        return district


def convertStdToAmount(standard):
    if standard >= 1 and standard <= 12:
        return 600
    if standard >= 13 and standard <= 17:
        return 2000
    return 0


def convertNumToStd(standard):
    """
    Parameter: Student Standard / Class Number
    Returns: String value if Num (eg: 11 = "+1", 13 = "1 DC")
    """
    std_data = {
        11: "+1",
        12: "+2",
        13: "1 DC",
        14: "2 DC",
        15: "3 DC",
        16: "1 PG",
        17: "2 PG",
    }
    if standard in std_data.keys():
        return std_data[standard]
    else:
        return standard


def convertParagraphToLine(text):
    """
    Parameters: String with newline breaks
    Returns: String without newline breaks
    """
    if "\n" in text:
        text = text.split("\n")
        text_stripped = []
        for line in text:
            if line != "":
                text_stripped.append(line.strip())
        # Join the list into a single string
        text_joined = " ".join(text_stripped)
        text = text_joined
    return text
