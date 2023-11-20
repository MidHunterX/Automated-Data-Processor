import pdfplumber   # PDF parsing
import datetime     # ISO Date format
import shutil       # Copying and Moving files
import docx         # Docx parsing
import csv          # CSV file manipulation
import os           # Directory path support
import sys          # Command line arguments and exit
import glob         # Finding files with extensions
import sqlite3      # SQLite DB operations
import pyperclip    # Clipboard handling
from pandas import DataFrame        # Printing Tables
from collections import Counter     # Most Common Value
from prompt_toolkit import prompt   # Prompt for Autocompletion
from sqlite3 import IntegrityError  # SQLite AccNo error
from prompt_toolkit.completion import WordCompleter  # Autocompletion Engine


def main():

    # ------------------------------------------------- [ INIT FILES AND DIRS ]

    input_dir = "input"
    ifsc_dataset = loadIfscDataset("data\\IFSC.csv")

    # ------------------------------------------------------------ [ COMMANDS ]

    command = getArgument()

    # Command List
    cmd_db = "database"
    cmd_form = "forms"
    cmd_ifsc = "ifsc"

    # PROCESS BRANCH
    if command == cmd_ifsc:
        getBranchFromPastedIfsc(ifsc_dataset)
        sys.exit(0)

    district_user = getDistrictFromUser()

    # Form Processing Variables
    if command == cmd_form:
        investigation_dir = initNestedDir(input_dir, "for checking")
        formatting_dir = initNestedDir(input_dir, "formatting issues")

    # Database Processing Variables
    elif command == cmd_db:
        db_file = "data\\database.db"
        iso_date = datetime.date.today().isoformat()
        input_dir = initNestedDir(input_dir, district_user)
        output_dir = initNestedDir(input_dir, iso_date)
        investigation_dir = initNestedDir(input_dir, "for checking")
        formatting_dir = initNestedDir(input_dir, "formatting issues")
        rejected_dir = initNestedDir(input_dir, "rejected")

    else:
        print("Unrecognized Command 😕")
        print(f"Try using: process {cmd_form} or process {cmd_db}")

    # ----------------------------------------------------- [ VARS FOR REPORT ]

    files_written = 0
    for_checking_count = 0
    incorrect_format_count = 0
    if command == cmd_db:
        rejected_count = 0

    # ----------------------------------------------------- [ FILE PROCESSING ]

    if command == "database":
        # Open connection to Database
        print("Connecting to Database")
        conn = sqlite3.connect(db_file)

    file_list = getFileList(input_dir, [".docx", ".pdf"])
    for file in file_list:

        printFileNameHeader(file)

        # Flags
        proceed = False
        valid_std = False

        # ---------------------------------------------------- [ FORM PARSING ]

        if correctFormat(file):
            proceed = True
            institution = getInstitutionDetails(file)
            student_data = getStudentDetails(file)
        else:
            print("⚠️ Formatting error detected!")

        # ------------------------------------------------- [ DATA PROCESSING ]

        if proceed is True:

            # Cleaning up Student Data for processing
            student_data = cleanStudentData(student_data)

            # Guessing District
            ifsc_list = getStudentIfscList(student_data)
            district_guess = guessDistrictFromIfscList(ifsc_list, ifsc_dataset)
            print(f"💡 Possible District: {district_guess}")

            # Deciding User District vs Guessed District
            district = district_user
            if district == "Unknown":
                if command == cmd_form:
                    district = district_guess
                if command == cmd_db:
                    district = getIndianState()
            print(f"✍️ Selected District: {district}\n")

            # Normalizing Student Data
            student_data = normalizeStudentData(student_data, ifsc_dataset)

            # Printing Final Data
            printInstitution(institution)
            print("")
            printStudentDataFrame(student_data, ifsc_dataset)

            # ------------------------------------------------ [ VERIFICATION ]

            # [ INFO: SAFE SPACE FOR KEYBOARDINTERRUPT ] #
            try:

                valid_std = isValidStudentStd(student_data)
                print("")
                # Enter to Confirm only if Student Class Valid
                if valid_std is True:
                    verification = input("Correct? (ret / n): ")
                else:
                    # Moves to investigation_dir for checking
                    print("⚠️ Cannot convert std to num equivalents :(")
                    verification = "n"

            # Abrupt ending for tactical retreat purposes (ctrl+c)
            except KeyboardInterrupt:

                print("Caught the Keyboard Interrupt ;D")
                if command == cmd_db:
                    # Close Connection to Database
                    print("Closing DB")
                    conn.close()
                break

            # ------------------------------------------- [ POST VERIFICATION ]

            if verification == "":
                print("✅ Marking as Correct.")

                # SORTING VERIFIED FORM
                if command == cmd_form:
                    # Creating district directory
                    output_dir = initNestedDir(input_dir, district)
                    shutil.move(file, output_dir)
                    files_written += 1

                # WRITING VERIFIED DATA INTO DATABASE
                if command == cmd_db:
                    if writeToDB(conn, district, institution, student_data):
                        print("✅ Data Written Successfully!")
                        shutil.move(file, output_dir)
                        files_written += 1
                    else:
                        print("❌ Rejected by Database")
                        shutil.move(file, rejected_dir)
                        rejected_count += 1
            else:
                # Enter for Further Investigation after rejecting verification
                try:
                    verification = input("Move for Investigation? (ret) ")
                    if verification == "":
                        print("❌ Moving for further Investigation.")
                        shutil.move(file, investigation_dir)
                        for_checking_count += 1
                # Abrupt ending for tactical retreat purposes (ctrl+c)
                except KeyboardInterrupt:
                    print("Caught the Keyboard Interrupt ;D")
                    break
        else:
            # Enter to Confirm after document formatting issues encountered
            try:
                verification = input("Proceed? (ret) ")
                if verification == "":
                    shutil.move(file, formatting_dir)
                incorrect_format_count += 1
            # Abrupt ending for tactical retreat purposes (ctrl+c)
            except KeyboardInterrupt:
                print("Caught the Keyboard Interrupt ;D")
                break

    if command == cmd_db:
        # Close Connection to Database
        print("Closing DB")
        conn.close()

    # -------------------------------------------------------------- [ REPORT ]

    print("")
    horizontal_line = "-"*80
    print(horizontal_line)
    print("FINAL REPORT".center(80))
    print(horizontal_line)
    print(f"Files Accepted    : {files_written}".center(80))
    print(f"For Checking      : {for_checking_count}".center(80))
    print(f"Formatting Issues : {incorrect_format_count}".center(80))
    if command == cmd_db:
        print(f"Rejected by DB    : {rejected_count}".center(80))
    print(horizontal_line)


# ================================ FUNCTIONS ================================ #


def loadDistrictDataset():
    district_list = [
        "Thiruvananthapuram", "Kollam", "Pathanamthitta", "Alappuzha",
        "Kottayam", "Idukki", "Ernakulam", "Thrissur", "Palakkad",
        "Malappuram", "Kozhikode", "Wayanad", "Kannur", "Kasargod"
    ]
    return district_list


def initNestedDir(input_dir, nest_name):
    directory_path = os.path.join(input_dir, nest_name)
    if not os.path.exists(directory_path):
        os.mkdir(directory_path)
    return directory_path


def getArgument():
    """
    Returns: A single argument from commandline
    """
    max_args = 1
    arg = sys.argv[1:]  # 0th arg is filename.py
    if len(arg) == max_args:
        command = str(sys.argv[1])
    elif len(arg) > max_args:
        sys.exit("Error: Too much arguments")
    elif len(arg) < 1:
        sys.exit("Error: No argument")
    return command


def correctPdfFormat(pdf_file):
    """
    Returns True if PDF is in correct Format
    """
    flags = {
        "Institution Heading": False,
        "Institution Lines": False,
        "Student Heading": False,
        "Student Table": False
    }
    try:
        with pdfplumber.open(pdf_file) as pdf:
            for page in pdf.pages:

                # ====== TEXT PARAGRAPH STARTS ====== #
                text = page.extract_text()

                # Check Heading: Institution Details
                if "Institution Details" in text:
                    flags["Institution Heading"] = True

                    # Check Length: Institution Details
                    start = text.index("Name of the Institution")
                    end = text.index("Student Details")
                    institution_details = text[start:end].splitlines()
                    if len(institution_details) == 4:
                        flags["Institution Lines"] = True

                # Check Heading: Student Details
                if "Student Details" in text:
                    flags["Student Heading"] = True

                # =========== TABLE STARTS =========== #
                table = page.extract_table()

                # Check Content: Student Table
                if table:
                    flags["Student Table"] = True
    except ValueError:
        pass

    # Logs
    if flags["Institution Heading"] is False:
        print("Heading not found: Institution Details")
    if flags["Institution Lines"] is False:
        print("Data Incomplete: Institution Details")
    if flags["Student Heading"] is False:
        print("Heading not found: Student Details")
    if flags["Student Table"] is False:
        print("Object not found: Student Table")

    status = all(flags.values())
    return status


def correctDocxFormat(docx_file):
    """
    Returns True if DOCX is in correct Format
    """
    inside_institution_details = False
    flags = {
        "name": False,
        "place": False,
        "number": False,
        "email": False,
    }
    doc = docx.Document(docx_file)
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if text.startswith("Institution Details"):
            inside_institution_details = True
        elif inside_institution_details:
            if text.startswith("Name of the Institution"):
                flags["name"] = True
            if text.startswith("Place"):
                flags["place"] = True
            if text.startswith("Phone number"):
                flags["number"] = True
            if text.startswith("Email Id"):
                flags["email"] = True

    # Logs
    if flags["name"] is False:
        print("Heading not found: Name of the Institution")
    if flags["place"] is False:
        print("Entry not found: Place")
    if flags["number"] is False:
        print("Entry not found: Number")
    if flags["email"] is False:
        print("Entry not found: Email")

    status = all(flags.values())
    return status


def getStudentDetailsPdf(pdf_file):
    """
    Parameter: PDF File
    Returns: A dictionary of tuples with Student details

    data = {
        0: (name, standard, ifsc, acc_no, holder, branch),
        1: (name, standard, ifsc, acc_no, holder, branch),
        2: (name, standard, ifsc, acc_no, holder, branch)
    }
    """
    with pdfplumber.open(pdf_file) as pdf:
        data = {}
        i = -1
        for page in pdf.pages:
            # Generate CSV list from PDF table
            table = page.extract_table()
            if table:
                for row in table:
                    # Replace \n substring with space
                    cleaned_row = []
                    for cell in row:
                        if isinstance(cell, str):
                            cleaned_row.append(cell.replace('\n', ' '))
                        else:
                            cleaned_row.append(cell)

                    name = cleaned_row[0]
                    standard = cleaned_row[1]
                    ifsc = cleaned_row[2]
                    acc_no = cleaned_row[3]
                    holder = cleaned_row[4]
                    branch = cleaned_row[5]

                    # Extracted data
                    if name:  # For avoiding empty rows
                        data[i] = name, standard, ifsc, acc_no, holder, branch
                        i = i + 1

    # Removes unwanted Header data
    data.pop(-1)

    return data


def getInstitutionDetailsPdf(pdf_file):
    """
    Parameters: Document.pdf file
    Returns: Dictionary of Institution Details

    data = {
        "name": name_of_institution,
        "place": place,
        "number": phone_number,
        "email": email_id
    }
    """

    with pdfplumber.open(pdf_file) as pdf:
        for page in pdf.pages:
            # Extract PDF as text
            text = page.extract_text()
            if "Institution Details" in text:
                start = text.index("Name of the Institution")
                end = text.index("Student Details")
                institution_details = text[start:end]

    # Splitting text at '\n' into a list
    lines = institution_details.split('\n')

    for line in lines:
        parts = line.split(':')
        if len(parts) == 2:
            key = parts[0].strip()
            value = parts[1].strip()

            # Assign values to variables
            if key == "Name of the Institution":
                name_of_institution = value
            elif key == "Place":
                place = value
            elif key == "Phone number":
                phone_number = value
            elif key == "Email Id":
                email_id = value

    # Extracted data
    data = {
        "name": name_of_institution,
        "place": place,
        "number": phone_number,
        "email": email_id
    }
    return data


def getFileList(dir, extensions):
    """
    Parameters: (dir, extensions)
        - dir: Directory Path
        - extensions: List of File extensions
    Returns: A list of file path.

    file_list = [file1.ext1, file2.ext1, file3.ext2, file4.ext2]
    """

    file_list = []

    if not isinstance(extensions, list):
        ext = [extensions]

    for ext in extensions:
        supported_files = glob.glob(os.path.join(dir, f"*{ext}"))
        for file in supported_files:
            file_list = file_list + [file]

    return file_list


def getInstitutionDetailsDocx(docx_file):
    """
    Parameters: Document.docx file
    Returns: Dictionary of Institution Details

    data = {
        "name": name_of_institution,
        "place": place,
        "number": phone_number,
        "email": email_id
    }
    """

    doc = docx.Document(docx_file)
    inside_institution_details = False
    name_of_institution = ""
    place = ""
    phone_number = ""
    email_id = ""

    for paragraph in doc.paragraphs:
        text = paragraph.text
        # Check if paragraph contains the "Institution Details"
        if text.startswith("Institution Details"):
            inside_institution_details = True
        elif inside_institution_details:
            # Split the paragraph by the colon
            parts = text.split(':')
            if len(parts) == 2:
                key = parts[0].strip()
                value = parts[1].strip()

                # Assign values to variables
                if key == "Name of the Institution":
                    name_of_institution = value
                elif key == "Place":
                    place = value
                elif key == "Phone number":
                    phone_number = value
                elif key == "Email Id":
                    email_id = value

    # Extracted data
    data = {
        "name": name_of_institution,
        "place": place,
        "number": phone_number,
        "email": email_id
    }
    return data


def getStudentDetailsDocx(docx_file):
    """
    Parameter: Document.docx file
    Returns: A dictionary of tuples with Student details

    data = {
        0: (name, standard, ifsc, acc_no, holder, branch),
        1: (name, standard, ifsc, acc_no, holder, branch),
        2: (name, standard, ifsc, acc_no, holder, branch)
    }
    """

    doc = docx.Document(docx_file)
    data = {}
    i = 0
    # Iterate through the tables in the document
    for table in doc.tables:
        for row in table.rows:
            first_column = row.cells[0].text
            if first_column != "" and first_column != "STUDENT NAME":
                name = row.cells[0].text
                standard = row.cells[1].text
                ifsc = row.cells[2].text
                acc_no = row.cells[3].text
                holder = row.cells[4].text
                branch = row.cells[5].text

                # Extracted data
                data[i] = name, standard, ifsc, acc_no, holder, branch
                i = i + 1
    return data


def getInstitutionDetails(file):
    """
    Parameters: Supported File
    Returns: Dictionary of Institution Details

    data = {
        "name": name_of_institution,
        "place": place,
        "number": phone_number,
        "email": email_id
    }
    """
    file_name, file_extension = os.path.basename(file).split(".")

    if file_extension == "docx":
        docx_file = file
        data = getInstitutionDetailsDocx(docx_file)

    if file_extension == "pdf":
        pdf_file = file
        data = getInstitutionDetailsPdf(pdf_file)

    return data


def getStudentDetails(file):
    """
    Parameter: Supported File
    Returns: A dictionary of tuples with Student details

    data = {
        0: (name, standard, ifsc, acc_no, holder, branch),
        1: (name, standard, ifsc, acc_no, holder, branch),
        2: (name, standard, ifsc, acc_no, holder, branch)
    }
    """
    file_name, file_extension = os.path.basename(file).split(".")

    if file_extension == "docx":
        docx_file = file
        data = getStudentDetailsDocx(docx_file)

    if file_extension == "pdf":
        pdf_file = file
        data = getStudentDetailsPdf(pdf_file)

    return data


def correctFormat(file):
    """
    Returns True if file is in correct Format
    """
    try:
        file_name, file_extension = os.path.basename(file).split(".")

        if file_extension == "docx":
            docx_file = file
            data = correctDocxFormat(docx_file)

        if file_extension == "pdf":
            pdf_file = file
            data = correctPdfFormat(pdf_file)

    except ValueError:
        print("⚠️ ValueError: Possible dot in file name")
        data = False

    return data


def getIndianState():
    indian_states = [
        "Andhra Pradesh",
        "Arunachal Pradesh",
        "Assam",
        "Bihar",
        "Chhattisgarh",
        "Goa",
        "Gujarat",
        "Haryana",
        "Himachal Pradesh",
        "Jharkhand",
        "Karnataka",
        "Kerala",
        "Madhya Pradesh",
        "Maharashtra",
        "Manipur",
        "Meghalaya",
        "Mizoram",
        "Nagaland",
        "Odisha",
        "Punjab",
        "Rajasthan",
        "Sikkim",
        "Tamil Nadu",
        "Telangana",
        "Tripura",
        "Uttar Pradesh",
        "Uttarakhand",
        "West Bengal"
    ]

    state_completer = WordCompleter(indian_states)
    # Prompt for state selection
    selected_state = prompt('Enter a state: ', completer=state_completer)
    return selected_state


def getDistrictFromUser():
    try:
        print("""
        1: TVM    6: IDK   11: KKD
        2: KLM    7: EKM   12: WYD
        3: PTA    8: TSR   13: KNR
        4: ALP    9: PKD   14: KSD
        5: KTM   10: MLP    0: Unknown
        """)
        district_dataset = loadDistrictDataset()
        district = "Unknown"
        data = int(input("Enter District No: "))
        data -= 1
        if data <= 13 and data >= 0:
            district = district_dataset[data]
        return district

    except ValueError:
        return district


# =============================== PROCEDURES ================================ #


def printTextBox_Centered(text):
    box_width = 80

    if len(text) > box_width - 4:
        text = text[:box_width - 7] + '...'

    freespace = box_width - len(text) - 4
    space = ' ' * (freespace // 2)

    horizontal_line = '+' + '-' * (box_width - 2) + '+'

    print(horizontal_line)
    print(f"|{space} {text}{space + ' ' if freespace % 2 != 0 else space}|")
    # print(f"|{' ' * (box_width - 2)}|")
    print(horizontal_line)


def printFileNameHeader(file):
    try:
        file_name, file_extension = os.path.basename(file).split(".")
        print("")
        printTextBox_Centered(f"📄 {file_name}.{file_extension}")
    except ValueError:
        printTextBox_Centered(f"📄 {file}")


def writeToCSV(csv_file, institution, student_data):
    """
    Parameter: (csv_file, institution, student_data)
    Returns: CSV File in working directory
    """
    with open(csv_file, mode="a", newline="") as file:
        writer = csv.writer(file)
        # Institution details
        for value in institution.values():
            writer.writerow([value])
        # Student details
        for row in student_data.values():
            writer.writerow(row)


def printInstitution(institution):
    inst_name = institution["name"]
    inst_place = institution["place"]
    inst_number = institution["number"]
    inst_email = institution["email"]
    print(f"{inst_name}\n{inst_place}\n{inst_number}\n{inst_email}")


def printStudentData(student_data):
    i = 1
    for key, value in student_data.items():
        name = value[0]
        standard = value[1]
        ifsc = value[2]
        acc_no = value[3]
        holder = value[4]
        branch = value[5]
        print(f"{i}: {name},{standard},{ifsc},{acc_no},{holder},{branch}")
        i += 1


def printStudentDataFrame(student_data, ifsc_dataset):
    student_record = []
    for key, val in student_data.items():

        ifsc = val[2]
        branch = getBranchFromIfsc(ifsc, ifsc_dataset)

        # Check IFSC Validity
        if branch == "":
            ifsc = f"{ifsc}❌"
        else:
            ifsc = f"{ifsc}✅"

        row = [key+1, val[0], val[1], ifsc, val[3], val[4], val[5]]
        student_record.append(row)

    df = DataFrame(
        student_record,
        columns=['', 'Name', 'Std', 'IFSC', 'Account No', 'Holder', 'Branch']
    )
    print(df.to_string(index=False))


def preprocessFiles(input_dir):
    """
    Renames files into numbers and
    Moves Unsupported files into a separate directory
    """
    unsupported_dir = os.path.join(input_dir, "unsupported")
    counter = 1

    # Ensure the unsupported directory exists
    if not os.path.exists(unsupported_dir):
        os.makedirs(unsupported_dir)

    for filename in os.listdir(input_dir):
        file_path = os.path.join(input_dir, filename)

        if os.path.isfile(file_path):
            # Check if it's a PDF or DOCX file
            if filename.lower().endswith(('.pdf', '.docx')):
                base_extension = os.path.splitext(filename)[1]
                new_name = f"{counter:03d}{base_extension}"
                new_path = os.path.join(input_dir, new_name)
                os.rename(file_path, new_path)
                counter += 1
            else:
                # Move unsupported files to the 'unsupported' directory
                unsupported_path = os.path.join(unsupported_dir, filename)
                shutil.move(file_path, unsupported_path)


# ======================= DATA PROCESSING FUNCTIONS ======================== #

# ------------------------------------------------------ [ DISTRICT FROM IFSC ]

def getStudentIfscList(student_data):
    """
    Parameter: Student Data from getStudentDetails()
    Returns: Iterable List of IFSC code

    ifsc = ["ifsc1", "ifsc2", "ifsc3", "ifsc4"]
    """
    i = 0
    ifsc = []
    for key, value in student_data.items():
        ifsc.append(value[2])
        i += 1
    return ifsc


def loadIfscDataset(csv_file):
    """
    Parameter: CSV Dataset from RazorPay
    Returns: Dataset Dictionary loaded into memory

    dataset[row['IFSC']] = {
        'Bank': row['BANK'],
        'Branch': row['BRANCH'],
        'Centre': row['CENTRE'],
        'District': row['DISTRICT'],
        'State': row['STATE'],
        'Address': row['ADDRESS'],
        'City': row['CITY'],
    }
    """
    dataset = {}
    with open(csv_file, mode='r', encoding='utf-8') as file:
        reader = csv.DictReader(file)
        for row in reader:
            dataset[row['IFSC']] = {
                'Bank': row['BANK'],
                'Branch': row['BRANCH'],
                'Centre': row['CENTRE'],
                'District': row['DISTRICT'],
                'State': row['STATE'],
                'Address': row['ADDRESS'],
                'City': row['CITY'],
            }
    return dataset


def getDistrictFromIfsc(ifsc, ifsc_dataset):
    """
    Parameters: (ifsc_code, ifsc_dataset)
    Returns: District as a String
    """
    district = "Unknown"
    ifsc_info = ifsc_dataset.get(ifsc)
    district_list = loadDistrictDataset()

    if ifsc_info:

        # District Finder (Initial Algorithm)
        district = ifsc_info["District"]
        # Return correct item from district data
        for item in district_list:
            if item.lower() == district.lower():
                district = item

        # District Finder v2.0 (New Algorithm)
        if district not in district_list:
            value_list = []
            for value in ifsc_info.values():
                value_list.append(value)
            district = get_most_common_value(value_list)

            # Return correct item from district data
            for item in district_list:
                if item.lower() == district.lower():
                    district = item

        # District Finder v1.0 (Fallback)
        if district not in district_list:
            for item in district_list:
                address = ifsc_info["Address"]
                if item.lower() in address.lower():
                    district = item

            # Return correct item from district data
            for item in district_list:
                if item.lower() == district.lower():
                    district = item

    return district


def get_most_common_value(a_list):
    """
    Parameters: A List of Values
    Returns: The most common value from list
    """
    count = Counter(a_list)
    mostCommon = count.most_common(1)
    return mostCommon[0][0]


def guessDistrictFromIfscList(ifsc_list, ifsc_dataset):
    """
    Parameters: (ifsc_list, ifsc_dataset)
    Returns: Guessed District as a String
    """
    district_list = []
    # Create a list of Districts
    for ifsc in ifsc_list:
        district = getDistrictFromIfsc(ifsc, ifsc_dataset)
        district_list.append(district)

    # Finding the most occured District
    district = get_most_common_value(district_list)

    # If District is Unknown, log data for user to verify
    if district == "Unknown":
        print("Couldn't decide District :(")
        print(f"Guess Data: {district_list}\n")

    # If District is unrecognized district, select from most common 3 values
    district_dataset = loadDistrictDataset()
    if district not in district_dataset:
        count = Counter(district_list)
        mostCommon_list = count.most_common(3)
        for i in mostCommon_list:
            if i[0] in district_dataset:
                district = i[0]

    return district


# ------------------------------------------------- [ CLASS NUMBER CONVERSION ]


def convertStdToNum(data):
    """
    Parameter: Student Standard / Class Number
    Returns: Numeric Value if String
    """
    data = str(data)
    data = data.strip().lower()

    std_dataset = {
        1: [
            "1",
            "1a",
            "1b",
            "1c",
            "1d",
            "1e",
            "1 a",
            "1 b",
            "1 c",
            "1 d",
            "1 e",
            "i",
            "1st",
            "one",
            "first"
        ],
        2: [
            "2",
            "2a",
            "2b",
            "2c",
            "2d",
            "2e",
            "2 a",
            "2 b",
            "2 c",
            "2 d",
            "2 e",
            "ii",
            "2nd",
            "two",
            "second"
        ],
        3: [
            "3",
            "3a",
            "3b",
            "3c",
            "3d",
            "3e",
            "3 a",
            "3 b",
            "3 c",
            "3 d",
            "3 e",
            "iii",
            "3rd",
            "three",
            "third"
        ],
        4: [
            "4",
            "4a",
            "4b",
            "4c",
            "4d",
            "4e",
            "4 a",
            "4 b",
            "4 c",
            "4 d",
            "4 e",
            "iv",
            "1v",
            "4th",
            "four",
            "fourth"
        ],
        5: [
            "5",
            "5a",
            "5b",
            "5c",
            "5d",
            "5e",
            "5 a",
            "5 b",
            "5 c",
            "5 d",
            "5 e",
            "v",
            "5th",
            "five",
            "fifth",
        ],
        6: [
            "6",
            "6a",
            "6b",
            "6c",
            "6d",
            "6e",
            "6 a",
            "6 b",
            "6 c",
            "6 d",
            "6 e",
            "vi",
            "v1",
            "six",
            "6th",
            "sixth",
        ],
        7: [
            "7",
            "7a",
            "7b",
            "7c",
            "7d",
            "7e",
            "7 a",
            "7 b",
            "7 c",
            "7 d",
            "7 e",
            "vii",
            "v11",
            "7th",
            "seven",
            "seventh",
        ],
        8: [
            "8",
            "8a",
            "8b",
            "8c",
            "8d",
            "8e",
            "8 a",
            "8 b",
            "8 c",
            "8 d",
            "8 e",
            "v111",
            "viii",
            "8th",
            "eight",
            "eighth",
        ],
        9: [
            "9",
            "9a",
            "9b",
            "9c",
            "9d",
            "9e",
            "9 a",
            "9 b",
            "9 c",
            "9 d",
            "9 e",
            "1x",
            "ix",
            "9th",
            "nine",
            "nineth",
        ],
        10: [
            "10",
            "10a",
            "10b",
            "10c",
            "10d",
            "10e",
            "10 a",
            "10 b",
            "10 c",
            "10 d",
            "10 e",
            "x",
            "10th",
            "ten",
            "tenth",
        ],
        11: [
            "11",
            "x1",
            "xi",
            "11th",
            "plus one",
            "plusone",
            "+1",
            "+1 science",
            "+1 commerce",
            "+1 humanities",
        ],
        12: [
            "12",
            "x11",
            "xii",
            "12th",
            "plus two",
            "plustwo",
            "+2",
            "+2 science",
            "+2 commerce",
            "+2 humanities",
        ],
        13: [
            "1 dc",
            "1dc",
            "i dc",
            "idc",
            "ist dc",
            "1stdc",
            "1st dc"
        ],
        14: [
            "2 dc",
            "2dc",
            "ii dc",
            "iidc",
            "iind dc",
            "2nddc",
            "2nd dc"
        ],
        15: [
            "3 dc",
            "3dc",
            "iii dc",
            "iiidc",
            "iiird dc",
            "3rddc",
            "3rd dc"
        ],
        16: [
            "1 pg",
            "1pg",
            "i pg",
            "ipg",
            "ist pg",
            "1st pg",
            "1stpg"
        ],
        17: [
            "2 pg",
            "2pg",
            "ii pg",
            "iipg",
            "iind pg",
            "2ndpg",
            "2nd pg"
        ],
    }
    if isinstance(data, str):
        data = data.lower()
        for key, values in std_dataset.items():
            for value in values:
                if data == value:
                    data = key
    return data


def convertParagraphToLine(text):
    """
    Parameters: String with newline breaks
    Returns: String without newline breaks
    """
    if "\n" in text:
        text = text.split("\n")
        text_stripped = []
        for line in text:
            if line != "":
                text_stripped.append(line.strip())
        # Join the list into a single string
        text_joined = " ".join(text_stripped)
        text = text_joined
    return text


def cleanStudentData(student_data):
    """
    Parameter: Student Data from getStudentDetails()
    Returns: A dictionary of tuples with cleaned up values for processing

    data = {
        0: (name, standard, ifsc, acc_no, holder, branch),
        1: (name, standard, ifsc, acc_no, holder, branch),
        2: (name, standard, ifsc, acc_no, holder, branch)
    }
    """
    i = 0
    data = {}
    for key, value in student_data.items():
        name = value[0]
        standard = value[1]
        ifsc = value[2]
        acc_no = value[3]
        holder = value[4]
        branch = value[5]

        # Cleaning spaces and newline from data
        name = convertParagraphToLine(name)
        standard = convertParagraphToLine(standard)
        holder = convertParagraphToLine(holder)
        branch = convertParagraphToLine(branch)

        # Cleaning up important data
        acc_no = str(acc_no).strip()
        ifsc = str(ifsc).strip()
        name = name.strip()
        holder = holder.strip()

        # Extracted data
        data[i] = name, standard, ifsc, acc_no, holder, branch
        i = i + 1

    return data


def getBranchFromIfsc(ifsc, ifsc_dataset):
    """
    Arguments: (ifsc, ifsc_dataset)
        - ifsc_list: IFSC Code
        - ifsc_dataset: IFSC Razorpay Dataset from loadIfscDataset()

    Returns:
        - "": If there exists no record of IFSC in dataset
        - branch: If Branch for IFSC is found
    """
    branch = ""
    if type(ifsc) is str:
        ifsc_details = ifsc_dataset.get(ifsc)
        if ifsc_details:
            branch = ifsc_details["Branch"]

    # Clean Up Branch
    if "IMPS" in branch:
        branch = branch.replace("IMPS", "").strip()

    return branch


def getBranchFromIfscList(ifsc_list, ifsc_dataset):
    """
    Arguments:
        - ifsc_list: List containing IFSC Codes
        - ifsc_dataset: IFSC Razorpay Dataset from loadIfscDataset()

    Returns:
        - branch_list: List of Branch for each IFSC Code
    """
    branch_list = []
    for ifsc in ifsc_list:
        branch = getBranchFromIfsc(ifsc, ifsc_dataset)
        branch_list.append(branch)
    return branch_list


def read_pasted_text():
    lines = []
    for line in sys.stdin:
        lines.append(line.rstrip('\n'))
    # Join the lines into a single string
    # pasted_text = '\n'.join(lines)
    return lines


def getBranchFromPastedIfsc(ifsc_dataset):

    print("")
    print("📝 Paste IFSC and press Ctrl+Z")
    print("-------------------------------")
    text = read_pasted_text()
    branch_list = []
    for ifsc in text:
        ifsc = ifsc.strip()
        branch = getBranchFromIfsc(ifsc, ifsc_dataset)
        branch_list.append(branch)

    # Join the list into a single string
    text = '\n'.join(branch_list)

    # Copy to Clipboard
    print("")
    print("✅ Copied to Clipboard")
    print("-----------------------")
    print(text)
    pyperclip.copy(text)


def normalizeStudentData(student_data, ifsc_dataset):
    """
    Parameter:
        - student_data: Student data from getStudentDetails()
        - ifsc_dataset: Razorpay IFSC Dataset from loadIfscDataset()

    Returns: A dictionary of tuples with corrected student data

    data = {
        0: (name, standard, ifsc, acc_no, holder, branch),
        1: (name, standard, ifsc, acc_no, holder, branch),
        2: (name, standard, ifsc, acc_no, holder, branch)
    }
    """
    i = 0
    data = {}
    for key, value in student_data.items():
        name = value[0]
        standard = value[1]
        ifsc = value[2]
        acc_no = value[3]
        holder = value[4]
        branch = value[5]

        # Normalizing Standard standard to Int variant
        standard = convertStdToNum(standard)

        # Normalizing Branch from IFSC using RazorPay Dataset
        rbi_branch = getBranchFromIfsc(ifsc, ifsc_dataset)

        # "," fix and long branch name fix
        if rbi_branch and "," not in rbi_branch and len(rbi_branch) < 20:
            branch = rbi_branch

        # Extracted data
        data[i] = name, standard, ifsc, acc_no, holder, branch
        i = i + 1

    return data


def isValidStudentStd(student_data):
    """
    Parameter: Student Data from getStudentDetails()
    Returns: Boolean Value (True / False)
    """
    data = []
    for key, value in student_data.items():
        standard = value[1]
        # Extracted data
        standard = convertStdToNum(standard)
        if type(standard) is int:
            data.append(True)
        else:
            data.append(False)

    return all(data)


def isValidIfsc(ifsc):
    """
    Parameter: IFSC Code
    Returns: True if correct and False if incorrect
    """
    ifsc = str(ifsc)
    # More than 5 characters
    if len(ifsc) < 6:
        return False
    # First 4 characters are alphabets (bank name)
    if not ifsc[:4].isalpha():
        return False
    # 5th character is 0 (for future use by RBI)
    if ifsc[4] != "0":
        return False
    # Return True if all checks are passed
    return True

# ========================== DATABASE OPERATIONS =========================== #


def writeToDB(conn, district, institution, student_data):
    """
    Arguments: (conn, district, institution, student_data)
        conn: Connection to database.db using sqlite3.connect()
        district: Name of District as String
        institution: Institution data from getInstitutionDetails function
        student_data: Student data from getStudentDetails function

    Returns:
        True: if inserted into DB successfully
        False: if any errors are encountered
    """

    try:
        cursor = conn.cursor()
        conn.execute("BEGIN TRANSACTION")

        # Insert Institution
        inst_name = institution["name"]
        inst_place = institution["place"]
        inst_number = institution["number"]
        inst_email = institution["email"]

        schoolSQL = """
        INSERT INTO Schools (
            SchoolName,
            District,
            Place,
            Phone,
            Email
        )
        VALUES ( ?, ?, ?, ?, ?)
        """
        values = inst_name, district, inst_place, inst_number, inst_email
        cursor.execute(schoolSQL, values)

        # Get the auto-incremented SchoolID
        school_id = cursor.lastrowid

        # Insert Students
        studentSQL = """
        INSERT INTO Students (
            SchoolID,
            StudentName,
            Class,
            IFSC,
            AccNo,
            AccHolder,
            Branch
        )
        VALUES ( ?, ?, ?, ?, ?, ?, ?)
        """

        for key, value in student_data.items():
            name = value[0]
            standard = value[1]
            ifsc = value[2]
            acc_no = value[3]
            holder = value[4]
            branch = value[5]
            variables = school_id, name, standard, ifsc, acc_no, holder, branch
            cursor.execute(studentSQL, variables)

        print("Commiting Changes")
        conn.commit()
        return True

    except IntegrityError as e:
        print(f"IntegrityError: {e}")
        conn.rollback()
        return False

    except Exception as e:
        print(f"Error: {e}")
        conn.rollback()
        return False


# ============================= MAIN FUNCTION ============================== #


if __name__ == "__main__":
    main()
