from django.shortcuts import render
from django.http import HttpResponse
from .forms import InstitutionForm, StudentFormSet
from .models import Institution


def generate_docx(request):
    if request.method == "POST":
        institution_form = InstitutionForm(request.POST)
        student_formset = StudentFormSet(request.POST, instance=Institution())

        if institution_form.is_valid() and student_formset.is_valid():
            # Save the institution details
            institution = institution_form.save()

            # Save the student details
            students = student_formset.save(commit=False)
            for student in students:
                student.institution = institution
                student.save()

            # Create a new docx document
            from docx import Document

            doc = Document()

            # Add institution details to the document
            doc.add_heading("Institution Details", level=1)
            doc.add_paragraph(f"Name of the Institution: {institution.name}")
            doc.add_paragraph(f"Place: {institution.place}")
            doc.add_paragraph(f"District: {institution.district}")
            doc.add_paragraph(f"Phone number: {institution.phone_no}")
            doc.add_paragraph(f"Email Id: {institution.email}")

            # Add student details to a table in the document
            doc.add_heading("Student Details", level=1)
            table = doc.add_table(rows=1, cols=6)
            table.style = "TableGrid"
            table.cell(0, 0).text = "STUDENT NAME"
            table.cell(0, 1).text = "CLASS"
            table.cell(0, 2).text = "IFSC CODE"
            table.cell(0, 3).text = "ACCOUNT NUMBER"
            table.cell(0, 4).text = "ACCOUNT HOLDER"
            table.cell(0, 5).text = "BRANCH"

            # Add each student to the table
            for student in students:
                row_cells = table.add_row().cells
                row_cells[0].text = student.student_name
                row_cells[1].text = student.student_class
                row_cells[2].text = student.student_ifsc
                row_cells[3].text = student.student_account
                row_cells[4].text = student.student_holder
                row_cells[5].text = student.student_branch

            # Save the document with the institution name as the filename
            filename = institution.name.replace(" ", "_") + ".docx"
            doc.save(filename)

            # Prepare the response to trigger download
            with open(filename, "rb") as file:
                response = HttpResponse(
                    file.read(),
                    content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
                response["Content-Disposition"] = f"attachment; filename={filename}"
            return response
    else:
        institution_form = InstitutionForm()
        student_formset = StudentFormSet(instance=Institution())

    return render(
        request,
        "generate_docx.html",
        {
            "institution_form": institution_form,
            "student_formset": student_formset},
    )
