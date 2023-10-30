import pdfplumber   # PDF parsing
import docx         # Docx parsing
import csv          # CSV file manipulation
import os           # Directory path support
import glob         # Finding files with extensions
import shutil       # Copying and Moving files
import datetime     # ISO Date
from collections import Counter     # Most Common Value


def main():

    # ------------------------------------------------- [ INIT FILES AND DIRS ]

    input_dir = "input"
    csv_file = "output.csv"
    iso_date = datetime.date.today().isoformat()
    output_dir = initNestedDir(input_dir, iso_date)
    investigation_dir = initNestedDir(input_dir, "for checking")
    ifsc_dataset = loadIfscDataset("data\\IFSC.csv")

    # ----------------------------------------------------- [ VARS FOR REPORT ]

    files_written = 0
    for_checking_count = 0
    incorrect_format_count = 0

    # ----------------------------------------------------- [ FILE PROCESSING ]

    preprocessFiles(input_dir)
    file_list = getFileList(input_dir, [".docx", ".pdf"])

    for file in file_list:
        file_name, file_extension = os.path.basename(file).split(".")
        print(file_name)
        print(file_extension)
        proceed = False

        if file_extension == "docx":
            docx_file = file
            if correctDocxFormat(docx_file):
                proceed = True
                print(f"\n==== {file_name}.{file_extension} ====")
                institution = getInstitutionDetailsDocx(docx_file)
                student_data = getStudentDetailsDocx(docx_file)

        if file_extension == "pdf":
            pdf_file = file
            if correctPdfFormat(pdf_file):
                proceed = True
                print(f"\n==== {file_name}.{file_extension} ====")
                institution = getInstitutionDetailsPdf(pdf_file)
                student_data = getStudentDetailsPdf(pdf_file)

        if proceed is True:
            ifsc_list = getStudentIfscList(student_data)
            district = guessDistrictFromIfscList(ifsc_list, ifsc_dataset)
            print(f"Possible District: {district}\n")

            student_data = normalizeStudentStd(student_data)

            printInstitution(institution)
            printStudentData(student_data)

            # Enter to Confirm
            verification = input("Correct? (ret / n): ")
            if verification == "":
                print("Marking as Correct.")
                writeToCSV(csv_file, institution, student_data)
                shutil.move(file, output_dir)
                files_written += 1
            else:
                print("Moving for further Investigation.")
                shutil.move(file, investigation_dir)
                for_checking_count += 1
        else:
            incorrect_format_count += 1

    # -------------------------------------------------------------- [ REPORT ]

    print("Final Report")
    print("------------")
    print(f"Files Written \t : {files_written}")
    print(f"For Checking \t : {for_checking_count}")
    print(f"Format Issues \t : {incorrect_format_count}")


# ================================ FUNCTIONS ================================ #


def initNestedDir(input_dir, nest_name):
    directory_path = os.path.join(input_dir, nest_name)
    if not os.path.exists(directory_path):
        os.mkdir(directory_path)
    return directory_path


def correctPdfFormat(pdf_file):
    """
    Returns True if PDF is in correct Format
    """
    flags = {
        "Institution Heading": False,
        "Institution Lines": False,
        "Student Heading": False,
        "Student Table": False
    }
    try:
        with pdfplumber.open(pdf_file) as pdf:
            for page in pdf.pages:

                # ====== TEXT PARAGRAPH STARTS ====== #
                text = page.extract_text()

                # Check Heading: Institution Details
                if "Institution Details" in text:
                    flags["Institution Heading"] = True

                    # Check Length: Institution Details
                    start = text.index("Name of the Institution")
                    end = text.index("Student Details")
                    institution_details = text[start:end].splitlines()
                    if len(institution_details) == 4:
                        flags["Institution Lines"] = True

                # Check Heading: Student Details
                if "Student Details" in text:
                    flags["Student Heading"] = True

                # =========== TABLE STARTS =========== #
                table = page.extract_table()

                # Check Content: Student Table
                if table:
                    flags["Student Table"] = True
    except ValueError:
        pass

    status = all(flags.values())
    return status


def correctDocxFormat(docx_file):
    """
    Returns True if DOCX is in correct Format
    """
    inside_institution_details = False
    flags = {
        "name": False,
        "place": False,
        "number": False,
        "email": False,
    }
    doc = docx.Document(docx_file)
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if text.startswith("Institution Details"):
            inside_institution_details = True
        elif inside_institution_details:
            if text.startswith("Name of the Institution"):
                flags["name"] = True
            if text.startswith("Place"):
                flags["place"] = True
            if text.startswith("Phone number"):
                flags["number"] = True
            if text.startswith("Email Id"):
                flags["email"] = True

    status = all(flags.values())
    return status


def getStudentDetailsPdf(pdf_file):
    """
    Parameter: PDF File
    Returns: A dictionary of tuples with Student details

    data = {
        0: (name, standard, ifsc, acc_no, holder, branch),
        1: (name, standard, ifsc, acc_no, holder, branch),
        2: (name, standard, ifsc, acc_no, holder, branch)
    }
    """
    with pdfplumber.open(pdf_file) as pdf:
        data = {}
        i = -1
        for page in pdf.pages:
            # Generate CSV list from PDF table
            table = page.extract_table()
            if table:
                for row in table:
                    # Replace \n substring with space
                    cleaned_row = []
                    for cell in row:
                        if isinstance(cell, str):
                            cleaned_row.append(cell.replace('\n', ' '))
                        else:
                            cleaned_row.append(cell)

                    name = cleaned_row[0]
                    standard = cleaned_row[1]
                    ifsc = cleaned_row[2]
                    acc_no = cleaned_row[3]
                    holder = cleaned_row[4]
                    branch = cleaned_row[5]

                    # Extracted data
                    if name:  # For avoiding empty rows
                        data[i] = name, standard, ifsc, acc_no, holder, branch
                        i = i + 1

    # Removes unwanted Header data
    data.pop(-1)

    return data


def getInstitutionDetailsPdf(pdf_file):
    """
    Parameters: Document.pdf file
    Returns: Dictionary of Institution Details

    data = {
        "name": name_of_institution,
        "place": place,
        "number": phone_number,
        "email": email_id
    }
    """

    with pdfplumber.open(pdf_file) as pdf:
        for page in pdf.pages:
            # Extract PDF as text
            text = page.extract_text()
            if "Institution Details" in text:
                start = text.index("Name of the Institution")
                end = text.index("Student Details")
                institution_details = text[start:end]

    # Splitting text at '\n' into a list
    lines = institution_details.split('\n')

    for line in lines:
        parts = line.split(':')
        if len(parts) == 2:
            key = parts[0].strip()
            value = parts[1].strip()

            # Assign values to variables
            if key == "Name of the Institution":
                name_of_institution = value
            elif key == "Place":
                place = value
            elif key == "Phone number":
                phone_number = value
            elif key == "Email Id":
                email_id = value

    # Extracted data
    data = {
        "name": name_of_institution,
        "place": place,
        "number": phone_number,
        "email": email_id
    }
    return data


def getFileList(dir, extensions):
    """
    Parameters: (dir, extensions)
        - dir: Directory Path
        - extensions: List of File extensions
    Returns: A list of file path.

    file_list = [file1.ext1, file2.ext1, file3.ext2, file4.ext2]
    """

    file_list = []

    if not isinstance(extensions, list):
        ext = [extensions]

    for ext in extensions:
        supported_files = glob.glob(os.path.join(dir, f"*{ext}"))
        for file in supported_files:
            file_list = file_list + [file]

    return file_list


def getPdfFileList(dir):
    """
    Parameters: Directory path
    Returns: A list of files with pdf extension.

    [file1.pdf, file2.pdf, file3.pdf]
    """

    pdf_files = glob.glob(os.path.join(dir, '*.pdf'))
    pdf_list = []
    for docx_file in pdf_files:
        pdf_list = pdf_list + [docx_file]
    return pdf_list


def getDocxFileList(dir):
    """
    Parameters: Directory path
    Returns: A list of files with docx extension.

    [file1.docx, file2.docx, file3.docx]
    """

    docx_files = glob.glob(os.path.join(dir, '*.docx'))
    docx_list = []
    for docx_file in docx_files:
        docx_list = docx_list + [docx_file]
    return docx_list


def getInstitutionDetailsDocx(docx_file):
    """
    Parameters: Document.docx file
    Returns: Dictionary of Institution Details

    data = {
        "name": name_of_institution,
        "place": place,
        "number": phone_number,
        "email": email_id
    }
    """

    doc = docx.Document(docx_file)
    inside_institution_details = False
    name_of_institution = ""
    place = ""
    phone_number = ""
    email_id = ""

    for paragraph in doc.paragraphs:
        text = paragraph.text
        # Check if paragraph contains the "Institution Details"
        if text.startswith("Institution Details"):
            inside_institution_details = True
        elif inside_institution_details:
            # Split the paragraph by the colon
            parts = text.split(':')
            if len(parts) == 2:
                key = parts[0].strip()
                value = parts[1].strip()

                # Assign values to variables
                if key == "Name of the Institution":
                    name_of_institution = value
                elif key == "Place":
                    place = value
                elif key == "Phone number":
                    phone_number = value
                elif key == "Email Id":
                    email_id = value

    # Extracted data
    data = {
        "name": name_of_institution,
        "place": place,
        "number": phone_number,
        "email": email_id
    }
    return data


def getStudentDetailsDocx(docx_file):
    """
    Parameter: Document.docx file
    Returns: A dictionary of tuples with Student details

    data = {
        0: (name, standard, ifsc, acc_no, holder, branch),
        1: (name, standard, ifsc, acc_no, holder, branch),
        2: (name, standard, ifsc, acc_no, holder, branch)
    }
    """

    doc = docx.Document(docx_file)
    data = {}
    i = 0
    # Iterate through the tables in the document
    for table in doc.tables:
        for row in table.rows:
            first_column = row.cells[0].text
            if first_column != "" and first_column != "STUDENT NAME":
                name = row.cells[0].text
                standard = row.cells[1].text
                ifsc = row.cells[2].text
                acc_no = row.cells[3].text
                holder = row.cells[4].text
                branch = row.cells[5].text

                # Extracted data
                data[i] = name, standard, ifsc, acc_no, holder, branch
                i = i + 1
    return data


# =============================== PROCEDURES ================================ #


def writeToCSV(csv_file, institution, student_data):
    """
    Parameter: (csv_file, institution, student_data)
    Returns: CSV File in working directory
    """
    with open(csv_file, mode="a", newline="") as file:
        writer = csv.writer(file)
        # Institution details
        for value in institution.values():
            writer.writerow([value])
        # Student details
        for row in student_data.values():
            writer.writerow(row)


def printInstitution(institution):
    inst_name = institution["name"]
    inst_place = institution["place"]
    inst_number = institution["number"]
    inst_email = institution["email"]
    print(f"{inst_name}\n{inst_place}\n{inst_number}\n{inst_email}")


def printStudentData(student_data):
    i = 1
    for key, value in student_data.items():
        name = value[0]
        standard = value[1]
        ifsc = value[2]
        acc_no = value[3]
        holder = value[4]
        branch = value[5]
        print(f"{i}: {name},{standard},{ifsc},{acc_no},{holder},{branch}")
        i += 1


def preprocessFiles(input_dir):
    """
    Renames files into numbers and
    Moves Unsupported files into a separate directory
    """
    unsupported_dir = os.path.join(input_dir, "unsupported")
    counter = 1

    # Ensure the unsupported directory exists
    if not os.path.exists(unsupported_dir):
        os.makedirs(unsupported_dir)

    for filename in os.listdir(input_dir):
        file_path = os.path.join(input_dir, filename)

        if os.path.isfile(file_path):
            # Check if it's a PDF or DOCX file
            if filename.lower().endswith(('.pdf', '.docx')):
                base_extension = os.path.splitext(filename)[1]
                new_name = f"{counter:03d}{base_extension}"
                new_path = os.path.join(input_dir, new_name)
                os.rename(file_path, new_path)
                counter += 1
            else:
                # Move unsupported files to the 'unsupported' directory
                unsupported_path = os.path.join(unsupported_dir, filename)
                shutil.move(file_path, unsupported_path)


# ======================= DATA PROCESSING FUNCTIONS ======================== #

# ------------------------------------------------------ [ DISTRICT FROM IFSC ]

def getStudentIfscList(student_data):
    """
    Parameter: Student Data from getStudentDetails()
    Returns: Iterable List of IFSC code

    ifsc = ["ifsc1", "ifsc2", "ifsc3", "ifsc4"]
    """
    i = 0
    ifsc = []
    for key, value in student_data.items():
        ifsc.append(value[2])
        i += 1
    return ifsc


def loadIfscDataset(csv_file):
    """
    Parameter: CSV Dataset from RazorPay
    Returns: Dataset Dictionary loaded into memory

    dataset[row['IFSC']] = {
        'Bank': row['BANK'],
        'Branch': row['BRANCH'],
        'Address': row['ADDRESS'],
        'District': row['DISTRICT'],
        'State': row['STATE']
    }
    """
    dataset = {}
    with open(csv_file, mode='r', encoding='utf-8') as file:
        reader = csv.DictReader(file)
        for row in reader:
            dataset[row['IFSC']] = {
                'Bank': row['BANK'],
                'Branch': row['BRANCH'],
                'Address': row['ADDRESS'],
                'District': row['DISTRICT'],
                'State': row['STATE']
            }
    return dataset


def getDistrictFromIfsc(ifsc, ifsc_dataset):
    """
    Parameters: (ifsc_code, ifsc_dataset)
    Returns: District as a String
    """
    district = "Unknown"
    ifsc_info = ifsc_dataset.get(ifsc)
    district_list = [
        "Alappuzha", "Ernakulam", "Idukki", "Kannur", "Kasargod",
        "Kollam", "Kottayam", "Kozhikode", "Malappuram", "Palakkad",
        "Pathanamthitta", "Thrissur", "Thiruvananthapuram", "Wayanad"
    ]
    if ifsc_info:
        district = ifsc_info["District"]

        # If unrecognized district, check address
        if district not in district_list:
            for item in district_list:
                address = ifsc_info["Address"]
                if item.lower() in address.lower():
                    district = item

        # Normalize district data
        for item in district_list:
            if item.lower() == district.lower():
                district = item

    return district


def get_most_common_value(a_list):
    """
    Parameters: A List of Values
    Returns: The most common value from list
    """
    count = Counter(a_list)
    mostCommon = count.most_common(1)
    return mostCommon[0][0]


def guessDistrictFromIfscList(ifsc_list, ifsc_dataset):
    """
    Parameters: (ifsc_list, ifsc_dataset)
    Returns: Guessed District as a String
    """
    district_list = []
    # Create a list of Districts
    for ifsc in ifsc_list:
        district = getDistrictFromIfsc(ifsc, ifsc_dataset)
        district_list.append(district)
    # Finding the most occured District
    return get_most_common_value(district_list)


# ------------------------------------------------- [ CLASS NUMBER CONVERSION ]


def convertStdToNum(data):
    """
    Parameter: Student Standard / Class Number
    Returns: Numeric Value if String
    """
    std_dataset = {
        1: ["1", "i", "1st", "first", "one"],
        2: ["2", "ii", "2nd", "second", "two"],
        3: ["3", "iii", "3rd", "third", "three"],
        4: ["4", "iv", "4th", "fourth", "four"],
        5: ["5", "v", "5th", "fifth", "five"],
        6: ["6", "vi", "6th", "sixth", "six"],
        7: ["7", "vii", "7th", "seventh", "seven"],
        8: ["8", "viii", "8th", "eighth", "eight"],
        9: ["9", "ix", "9th", "nineth", "nine"],
        10: ["10", "x", "10th", "tenth", "ten"],
        11: ["11", "xi", "11th", "plus one", "+1", "plusone"],
        12: ["12", "xii", "12th", "plus two", "+2", "plustwo"],
        13: ["1dc", "1 dc", "ist dc", "i dc", "idc", "1stdc", "1st dc"],
        14: ["2dc", "2 dc", "iind dc", "ii dc", "iidc", "2nddc", "2nd dc"],
        15: ["3dc", "3 dc", "iiird dc", "iii dc", "iiidc", "3rddc", "3rd dc"],
        16: ["1pg", "1 pg", "ist pg", "i pg", "ipg", "1stpg", "1st pg"],
        17: ["2pg", "2 pg", "iind pg", "ii pg", "iipg", "2ndpg", "2nd pg"],
    }
    if isinstance(data, str):
        data = data.lower()
        for key, values in std_dataset.items():
            for value in values:
                if data == value:
                    data = key
    return data


def normalizeStudentStd(student_data):
    """
    Parameter: Student Data from getStudentDetails()
    Returns: A dictionary of tuples with corrected Student standard

    data = {
        0: (name, numeric_standard, ifsc, acc_no, holder, branch),
        1: (name, numeric_standard, ifsc, acc_no, holder, branch),
        2: (name, numeric_standard, ifsc, acc_no, holder, branch)
    }
    """
    i = 0
    data = {}
    for key, value in student_data.items():
        name = value[0]
        standard = value[1]
        ifsc = value[2]
        acc_no = value[3]
        holder = value[4]
        branch = value[5]

        # Extracted data
        data[i] = name, convertStdToNum(standard), ifsc, acc_no, holder, branch
        i = i + 1

    return data


# ============================= MAIN FUNCTION ============================== #


if __name__ == "__main__":
    main()
