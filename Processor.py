from sqlite3 import IntegrityError
import pdfplumber   # PDF parsing
import docx         # Docx parsing
import csv          # CSV file manipulation
import os           # Directory path support
import glob         # Finding files with extensions
import shutil       # Copying and Moving files
import pandas       # Printing Tables
from collections import Counter     # Most Common Value


def main():

    # ------------------------------------------------- [ INIT FILES AND DIRS ]

    input_dir = "input"
    investigation_dir = initNestedDir(input_dir, "for checking")
    formatting_dir = initNestedDir(input_dir, "formatting issues")

    ifsc_dataset = loadIfscDataset("data\\IFSC.csv")
    district_user = getDistrictFromUser()

    # ----------------------------------------------------- [ VARS FOR REPORT ]

    files_written = 0
    for_checking_count = 0
    incorrect_format_count = 0

    # ----------------------------------------------------- [ FILE PROCESSING ]

    file_list = getFileList(input_dir, [".docx", ".pdf"])

    for file in file_list:
        file_name, file_extension = os.path.basename(file).split(".")

        # Flags
        proceed = False
        valid_std = False

        # Check for Formatting issues
        print(f"\n==== {file_name}.{file_extension} ====")

        if correctFormat(file):
            proceed = True
            institution = getInstitutionDetails(file)
            student_data = getStudentDetails(file)
        else:
            print("Formatting error detected!")

        # If Correct format, proceed
        if proceed is True:

            # Guessing District
            ifsc_list = getStudentIfscList(student_data)
            district_guess = guessDistrictFromIfscList(ifsc_list, ifsc_dataset)
            print(f"Possible District: {district_guess}")

            # Deciding User District vs Guessed District
            district = district_user
            if district == "Unknown":
                district = district_guess
            print(f"Selected District: {district}\n")

            # Validating Student Data for conversion to int equivalent
            if isValidStudentStd(student_data):
                valid_std = True
            else:
                print("Cannot convert all standards to num equivalents :(")

            # Normalizing Student Data
            student_data = normalizeStudentStd(student_data)
            student_data = normalizeStudentBranch(student_data, ifsc_dataset)

            # Printing Final Data
            printInstitution(institution)
            print("")
            printStudentDataFrame(student_data)

            # Enter to Confirm only if Student Class Valid
            if valid_std is True:
                verification = input("\nCorrect? (ret / n): ")
            else:
                # Moves to investigation_dir for checking
                verification = "n"

            if verification == "":
                print("Marking as Correct.")
                # Creating district directory
                output_dir = initNestedDir(input_dir, district)
                shutil.move(file, output_dir)
                files_written += 1
            else:
                print("Moving for further Investigation.")
                shutil.move(file, investigation_dir)
                for_checking_count += 1
        else:
            # Enter to Confirm
            verification = input("Proceed? (ret)")
            if verification == "":
                shutil.move(file, formatting_dir)
            incorrect_format_count += 1

    # -------------------------------------------------------------- [ REPORT ]

    print("")
    print("Final Report")
    print("------------")
    print(f"Files Accepted \t\t : {files_written}")
    print(f"For Checking \t\t : {for_checking_count}")
    print(f"Formatting Issues \t : {incorrect_format_count}")


# ================================ FUNCTIONS ================================ #


def loadDistrictDataset():
    district_list = [
        "Thiruvananthapuram", "Kollam", "Pathanamthitta", "Alappuzha",
        "Kottayam", "Idukki", "Ernakulam", "Thrissur", "Palakkad",
        "Malappuram", "Kozhikode", "Wayanad", "Kannur", "Kasargod"
    ]
    return district_list


def initNestedDir(input_dir, nest_name):
    directory_path = os.path.join(input_dir, nest_name)
    if not os.path.exists(directory_path):
        os.mkdir(directory_path)
    return directory_path


def correctPdfFormat(pdf_file):
    """
    Returns True if PDF is in correct Format
    """
    flags = {
        "Institution Heading": False,
        "Institution Lines": False,
        "Student Heading": False,
        "Student Table": False
    }
    try:
        with pdfplumber.open(pdf_file) as pdf:
            for page in pdf.pages:

                # ====== TEXT PARAGRAPH STARTS ====== #
                text = page.extract_text()

                # Check Heading: Institution Details
                if "Institution Details" in text:
                    flags["Institution Heading"] = True

                    # Check Length: Institution Details
                    start = text.index("Name of the Institution")
                    end = text.index("Student Details")
                    institution_details = text[start:end].splitlines()
                    if len(institution_details) == 4:
                        flags["Institution Lines"] = True

                # Check Heading: Student Details
                if "Student Details" in text:
                    flags["Student Heading"] = True

                # =========== TABLE STARTS =========== #
                table = page.extract_table()

                # Check Content: Student Table
                if table:
                    flags["Student Table"] = True
    except ValueError:
        pass

    status = all(flags.values())
    return status


def correctDocxFormat(docx_file):
    """
    Returns True if DOCX is in correct Format
    """
    inside_institution_details = False
    flags = {
        "name": False,
        "place": False,
        "number": False,
        "email": False,
    }
    doc = docx.Document(docx_file)
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if text.startswith("Institution Details"):
            inside_institution_details = True
        elif inside_institution_details:
            if text.startswith("Name of the Institution"):
                flags["name"] = True
            if text.startswith("Place"):
                flags["place"] = True
            if text.startswith("Phone number"):
                flags["number"] = True
            if text.startswith("Email Id"):
                flags["email"] = True

    # Logs
    if flags["name"] is False:
        print("Heading not found: Name of the Institution")
    if flags["place"] is False:
        print("Entry not found: Place")
    if flags["number"] is False:
        print("Entry not found: Number")
    if flags["email"] is False:
        print("Entry not found: Email")

    status = all(flags.values())
    return status


def getStudentDetailsPdf(pdf_file):
    """
    Parameter: PDF File
    Returns: A dictionary of tuples with Student details

    data = {
        0: (name, standard, ifsc, acc_no, holder, branch),
        1: (name, standard, ifsc, acc_no, holder, branch),
        2: (name, standard, ifsc, acc_no, holder, branch)
    }
    """
    with pdfplumber.open(pdf_file) as pdf:
        data = {}
        i = -1
        for page in pdf.pages:
            # Generate CSV list from PDF table
            table = page.extract_table()
            if table:
                for row in table:
                    # Replace \n substring with space
                    cleaned_row = []
                    for cell in row:
                        if isinstance(cell, str):
                            cleaned_row.append(cell.replace('\n', ' '))
                        else:
                            cleaned_row.append(cell)

                    name = cleaned_row[0]
                    standard = cleaned_row[1]
                    ifsc = cleaned_row[2]
                    acc_no = cleaned_row[3]
                    holder = cleaned_row[4]
                    branch = cleaned_row[5]

                    # Extracted data
                    if name:  # For avoiding empty rows
                        data[i] = name, standard, ifsc, acc_no, holder, branch
                        i = i + 1

    # Removes unwanted Header data
    data.pop(-1)

    return data


def getInstitutionDetailsPdf(pdf_file):
    """
    Parameters: Document.pdf file
    Returns: Dictionary of Institution Details

    data = {
        "name": name_of_institution,
        "place": place,
        "number": phone_number,
        "email": email_id
    }
    """

    with pdfplumber.open(pdf_file) as pdf:
        for page in pdf.pages:
            # Extract PDF as text
            text = page.extract_text()
            if "Institution Details" in text:
                start = text.index("Name of the Institution")
                end = text.index("Student Details")
                institution_details = text[start:end]

    # Splitting text at '\n' into a list
    lines = institution_details.split('\n')

    for line in lines:
        parts = line.split(':')
        if len(parts) == 2:
            key = parts[0].strip()
            value = parts[1].strip()

            # Assign values to variables
            if key == "Name of the Institution":
                name_of_institution = value
            elif key == "Place":
                place = value
            elif key == "Phone number":
                phone_number = value
            elif key == "Email Id":
                email_id = value

    # Extracted data
    data = {
        "name": name_of_institution,
        "place": place,
        "number": phone_number,
        "email": email_id
    }
    return data


def getFileList(dir, extensions):
    """
    Parameters: (dir, extensions)
        - dir: Directory Path
        - extensions: List of File extensions
    Returns: A list of file path.

    file_list = [file1.ext1, file2.ext1, file3.ext2, file4.ext2]
    """

    file_list = []

    if not isinstance(extensions, list):
        ext = [extensions]

    for ext in extensions:
        supported_files = glob.glob(os.path.join(dir, f"*{ext}"))
        for file in supported_files:
            file_list = file_list + [file]

    return file_list


def getInstitutionDetailsDocx(docx_file):
    """
    Parameters: Document.docx file
    Returns: Dictionary of Institution Details

    data = {
        "name": name_of_institution,
        "place": place,
        "number": phone_number,
        "email": email_id
    }
    """

    doc = docx.Document(docx_file)
    inside_institution_details = False
    name_of_institution = ""
    place = ""
    phone_number = ""
    email_id = ""

    for paragraph in doc.paragraphs:
        text = paragraph.text
        # Check if paragraph contains the "Institution Details"
        if text.startswith("Institution Details"):
            inside_institution_details = True
        elif inside_institution_details:
            # Split the paragraph by the colon
            parts = text.split(':')
            if len(parts) == 2:
                key = parts[0].strip()
                value = parts[1].strip()

                # Assign values to variables
                if key == "Name of the Institution":
                    name_of_institution = value
                elif key == "Place":
                    place = value
                elif key == "Phone number":
                    phone_number = value
                elif key == "Email Id":
                    email_id = value

    # Extracted data
    data = {
        "name": name_of_institution,
        "place": place,
        "number": phone_number,
        "email": email_id
    }
    return data


def getStudentDetailsDocx(docx_file):
    """
    Parameter: Document.docx file
    Returns: A dictionary of tuples with Student details

    data = {
        0: (name, standard, ifsc, acc_no, holder, branch),
        1: (name, standard, ifsc, acc_no, holder, branch),
        2: (name, standard, ifsc, acc_no, holder, branch)
    }
    """

    doc = docx.Document(docx_file)
    data = {}
    i = 0
    # Iterate through the tables in the document
    for table in doc.tables:
        for row in table.rows:
            first_column = row.cells[0].text
            if first_column != "" and first_column != "STUDENT NAME":
                name = row.cells[0].text
                standard = row.cells[1].text
                ifsc = row.cells[2].text
                acc_no = row.cells[3].text
                holder = row.cells[4].text
                branch = row.cells[5].text

                # Extracted data
                data[i] = name, standard, ifsc, acc_no, holder, branch
                i = i + 1
    return data


def getInstitutionDetails(file):
    """
    Parameters: Supported File
    Returns: Dictionary of Institution Details

    data = {
        "name": name_of_institution,
        "place": place,
        "number": phone_number,
        "email": email_id
    }
    """
    file_name, file_extension = os.path.basename(file).split(".")

    if file_extension == "docx":
        docx_file = file
        data = getInstitutionDetailsDocx(docx_file)

    if file_extension == "pdf":
        pdf_file = file
        data = getInstitutionDetailsPdf(pdf_file)

    return data


def getStudentDetails(file):
    """
    Parameter: Supported File
    Returns: A dictionary of tuples with Student details

    data = {
        0: (name, standard, ifsc, acc_no, holder, branch),
        1: (name, standard, ifsc, acc_no, holder, branch),
        2: (name, standard, ifsc, acc_no, holder, branch)
    }
    """
    file_name, file_extension = os.path.basename(file).split(".")

    if file_extension == "docx":
        docx_file = file
        data = getStudentDetailsDocx(docx_file)

    if file_extension == "pdf":
        pdf_file = file
        data = getStudentDetailsPdf(pdf_file)

    return data


def correctFormat(file):
    """
    Returns True if file is in correct Format
    """
    file_name, file_extension = os.path.basename(file).split(".")

    if file_extension == "docx":
        docx_file = file
        data = correctDocxFormat(docx_file)

    if file_extension == "pdf":
        pdf_file = file
        data = correctPdfFormat(pdf_file)

    return data


# =============================== PROCEDURES ================================ #


def getDistrictFromUser():
    try:
        print("""
        1: TVM    6: IDK   11: KKD
        2: KLM    7: EKM   12: WYD
        3: PTA    8: TSR   13: KNR
        4: ALP    9: PKD   14: KSD
        5: KTM   10: MLP    0: Unknown
        """)
        district_dataset = loadDistrictDataset()
        district = "Unknown"
        data = int(input("Enter District No: "))
        data -= 1
        if data <= 13 and data >= 0:
            district = district_dataset[data]
        return district

    except ValueError:
        return district


def writeToCSV(csv_file, institution, student_data):
    """
    Parameter: (csv_file, institution, student_data)
    Returns: CSV File in working directory
    """
    with open(csv_file, mode="a", newline="") as file:
        writer = csv.writer(file)
        # Institution details
        for value in institution.values():
            writer.writerow([value])
        # Student details
        for row in student_data.values():
            writer.writerow(row)


def printInstitution(institution):
    inst_name = institution["name"]
    inst_place = institution["place"]
    inst_number = institution["number"]
    inst_email = institution["email"]
    print(f"{inst_name}\n{inst_place}\n{inst_number}\n{inst_email}")


def printStudentData(student_data):
    i = 1
    for key, value in student_data.items():
        name = value[0]
        standard = value[1]
        ifsc = value[2]
        acc_no = value[3]
        holder = value[4]
        branch = value[5]
        print(f"{i}: {name},{standard},{ifsc},{acc_no},{holder},{branch}")
        i += 1


def printStudentDataFrame(student_data):
    student_record = []
    for key, val in student_data.items():
        row = [key+1, val[0], val[1], val[2], val[3], val[4], val[5]]
        student_record.append(row)

    df = pandas.DataFrame(
        student_record,
        columns=['', 'Name', 'Std', 'IFSC', 'Account No', 'Holder', 'Branch']
    )
    print(df.to_string(index=False))


def preprocessFiles(input_dir):
    """
    Renames files into numbers and
    Moves Unsupported files into a separate directory
    """
    unsupported_dir = os.path.join(input_dir, "unsupported")
    counter = 1

    # Ensure the unsupported directory exists
    if not os.path.exists(unsupported_dir):
        os.makedirs(unsupported_dir)

    for filename in os.listdir(input_dir):
        file_path = os.path.join(input_dir, filename)

        if os.path.isfile(file_path):
            # Check if it's a PDF or DOCX file
            if filename.lower().endswith(('.pdf', '.docx')):
                base_extension = os.path.splitext(filename)[1]
                new_name = f"{counter:03d}{base_extension}"
                new_path = os.path.join(input_dir, new_name)
                os.rename(file_path, new_path)
                counter += 1
            else:
                # Move unsupported files to the 'unsupported' directory
                unsupported_path = os.path.join(unsupported_dir, filename)
                shutil.move(file_path, unsupported_path)


# ======================= DATA PROCESSING FUNCTIONS ======================== #

# ------------------------------------------------------ [ DISTRICT FROM IFSC ]

def getStudentIfscList(student_data):
    """
    Parameter: Student Data from getStudentDetails()
    Returns: Iterable List of IFSC code

    ifsc = ["ifsc1", "ifsc2", "ifsc3", "ifsc4"]
    """
    i = 0
    ifsc = []
    for key, value in student_data.items():
        ifsc.append(value[2])
        i += 1
    return ifsc


def loadIfscDataset(csv_file):
    """
    Parameter: CSV Dataset from RazorPay
    Returns: Dataset Dictionary loaded into memory

    dataset[row['IFSC']] = {
        'Bank': row['BANK'],
        'Branch': row['BRANCH'],
        'Centre': row['CENTRE'],
        'District': row['DISTRICT'],
        'State': row['STATE'],
        'Address': row['ADDRESS'],
        'City': row['CITY'],
    }
    """
    dataset = {}
    with open(csv_file, mode='r', encoding='utf-8') as file:
        reader = csv.DictReader(file)
        for row in reader:
            dataset[row['IFSC']] = {
                'Bank': row['BANK'],
                'Branch': row['BRANCH'],
                'Centre': row['CENTRE'],
                'District': row['DISTRICT'],
                'State': row['STATE'],
                'Address': row['ADDRESS'],
                'City': row['CITY'],
            }
    return dataset


def getDistrictFromIfsc(ifsc, ifsc_dataset):
    """
    Parameters: (ifsc_code, ifsc_dataset)
    Returns: District as a String
    """
    district = "Unknown"
    ifsc_info = ifsc_dataset.get(ifsc)
    district_list = loadDistrictDataset()

    if ifsc_info:

        # District Finder (Initial Algorithm)
        district = ifsc_info["District"]
        # Return correct item from district data
        for item in district_list:
            if item.lower() == district.lower():
                district = item

        # District Finder v2.0 (New Algorithm)
        if district not in district_list:
            value_list = []
            for value in ifsc_info.values():
                value_list.append(value)
            district = get_most_common_value(value_list)

            # Return correct item from district data
            for item in district_list:
                if item.lower() == district.lower():
                    district = item

        # District Finder v1.0 (Fallback)
        if district not in district_list:
            for item in district_list:
                address = ifsc_info["Address"]
                if item.lower() in address.lower():
                    district = item

            # Return correct item from district data
            for item in district_list:
                if item.lower() == district.lower():
                    district = item

    return district


def get_most_common_value(a_list):
    """
    Parameters: A List of Values
    Returns: The most common value from list
    """
    count = Counter(a_list)
    mostCommon = count.most_common(1)
    return mostCommon[0][0]


def guessDistrictFromIfscList(ifsc_list, ifsc_dataset):
    """
    Parameters: (ifsc_list, ifsc_dataset)
    Returns: Guessed District as a String
    """
    district_list = []
    # Create a list of Districts
    for ifsc in ifsc_list:
        district = getDistrictFromIfsc(ifsc, ifsc_dataset)
        district_list.append(district)

    # Finding the most occured District
    district = get_most_common_value(district_list)

    # If District is Unknown, log data for user to verify
    if district == "Unknown":
        print("Couldn't decide District :(")
        print(f"Guess Data: {district_list}\n")

    return district


# ------------------------------------------------- [ CLASS NUMBER CONVERSION ]


def convertStdToNum(data):
    """
    Parameter: Student Standard / Class Number
    Returns: Numeric Value if String
    """
    data = str(data)
    data = data.strip().lower()

    std_dataset = {
        1: [
            "1", "one",
            "1a", "1b", "1c", "1d", "1e",
            "i",
            "1st", "first"
        ],
        2: [
            "2", "two"
            "2a", "2b", "2c", "2d", "2e",
            "ii",
            "2nd", "second"
        ],
        3: [
            "3", "three",
            "3a", "3b", "3c", "3d", "3e",
            "iii",
            "3rd", "third"
        ],
        4: [
            "4", "four",
            "4a", "4b", "4c", "4d", "4e",
            "1v", "iv",
            "4th", "fourth"
        ],
        5: [
            "5", "five",
            "5a", "5b", "5c", "5d", "5e",
            "v",
            "5th", "fifth",
        ],
        6: [
            "6", "six",
            "6a", "6b", "6c", "6d", "6e",
            "v1", "vi",
            "6th", "sixth",
        ],
        7: [
            "7", "seven",
            "7a", "7b", "7c", "7d", "7e",
            "v11", "vii",
            "7th", "seventh",
        ],
        8: [
            "8", "eight",
            "8a", "8b", "8c", "8d", "8e",
            "v111", "viii",
            "8th", "eighth",
        ],
        9: [
            "9", "nine",
            "9a", "9b", "9c", "9d", "9e",
            "1x", "ix",
            "9th", "nineth",
        ],
        10: [
            "10", "ten",
            "10a", "10b", "10c", "10d", "10e",
            "x",
            "10th", "tenth",
        ],
        11: [
            "11",
            "x1", "xi",
            "11th",
            "plus one", "plusone",
            "+1", "+1 science", "+1 commerce", "+1 humanities",
        ],
        12: [
            "12",
            "x11", "xii",
            "12th",
            "plus two", "plustwo",
            "+2", "+2 science", "+2 commerce", "+2 humanities",
        ],
        13: [
            "1 dc", "1dc",
            "i dc", "idc",
            "ist dc", "1stdc", "1st dc"
        ],
        14: [
            "2 dc", "2dc",
            "ii dc", "iidc",
            "iind dc", "2nddc", "2nd dc"
        ],
        15: [
            "3 dc", "3dc",
            "iii dc", "iiidc",
            "iiird dc", "3rddc", "3rd dc"
        ],
        16: [
            "1 pg", "1pg",
            "i pg", "ipg",
            "ist pg", "1st pg", "1stpg"
        ],
        17: [
            "2 pg", "2pg",
            "ii pg", "iipg",
            "iind pg", "2ndpg", "2nd pg"
        ],
    }
    if isinstance(data, str):
        data = data.lower()
        for key, values in std_dataset.items():
            for value in values:
                if data == value:
                    data = key
    return data


def convertParagraphToLine(text):
    """
    Parameters: String with newline breaks
    Returns: String without newline breaks
    """
    if "\n" in text:
        text = text.split("\n")
        text_stripped = []
        for line in text:
            if line != "":
                text_stripped.append(line.strip())
        # Join the list into a single string
        text_joined = " ".join(text_stripped)
        text = text_joined
    return text


def normalizeStudentStd(student_data):
    """
    Parameter: Student Data from getStudentDetails()
    Returns: A dictionary of tuples with corrected Student standard

    data = {
        0: (name, numeric_standard, ifsc, acc_no, holder, branch),
        1: (name, numeric_standard, ifsc, acc_no, holder, branch),
        2: (name, numeric_standard, ifsc, acc_no, holder, branch)
    }
    """
    i = 0
    data = {}
    for key, value in student_data.items():
        name = value[0]
        standard = value[1]
        ifsc = value[2]
        acc_no = value[3]
        holder = value[4]
        branch = value[5]

        # Cleaning spaces and newline from data
        name = convertParagraphToLine(name)
        standard = convertParagraphToLine(standard)
        holder = convertParagraphToLine(holder)
        branch = convertParagraphToLine(branch)

        # Cleaning up important data
        acc_no = str(acc_no).strip()
        ifsc = str(ifsc).strip()

        # Normalizing Standard data to Int variant
        standard = convertStdToNum(standard)

        # Extracted data
        data[i] = name, standard, ifsc, acc_no, holder, branch
        i = i + 1

    return data


def getBranchFromIfsc(ifsc, ifsc_dataset):
    """
    Arguments: (ifsc, ifsc_dataset)
        - ifsc_list: IFSC Code
        - ifsc_dataset: IFSC Razorpay Dataset from loadIfscDataset()

    Returns:
        - "": If there exists no record of IFSC in dataset
        - branch: If Branch for IFSC is found
    """
    branch = ""
    if type(ifsc) is str:
        ifsc_details = ifsc_dataset.get(ifsc)
        if ifsc_details:
            branch = ifsc_details["Branch"]
    return branch


def normalizeStudentBranch(student_data, ifsc_dataset):
    """
    Parameter:
        - student_data: Student data from getStudentDetails()
        - ifsc_dataset: Razorpay IFSC Dataset from loadIfscDataset()

    Returns: A dictionary of tuples with corrected Branch

    data = {
        0: (name, standard, ifsc, acc_no, holder, razorpay_branch),
        1: (name, standard, ifsc, acc_no, holder, razorpay_branch),
        2: (name, standard, ifsc, acc_no, holder, razorpay_branch)
    }
    """
    i = 0
    data = {}
    for key, value in student_data.items():
        name = value[0]
        standard = value[1]
        ifsc = value[2]
        acc_no = value[3]
        holder = value[4]
        branch = value[5]

        # Normalizing Branch from IFSC using RazorPay Dataset
        razorpay_branch = getBranchFromIfsc(ifsc, ifsc_dataset)
        if razorpay_branch:
            branch = razorpay_branch

        # Extracted data
        data[i] = name, standard, ifsc, acc_no, holder, branch
        i = i + 1

    return data


def isValidStudentStd(student_data):
    """
    Parameter: Student Data from getStudentDetails()
    Returns: Boolean Value (True / False)
    """
    data = []
    for key, value in student_data.items():
        standard = value[1]
        # Extracted data
        standard = convertStdToNum(standard)
        if type(standard) is int:
            data.append(True)
        else:
            data.append(False)

    return all(data)


# ========================== DATABASE OPERATIONS =========================== #


def writeToDB(conn, district, institution, student_data):
    """
    Arguments: (conn, district, institution, student_data)
        conn: Connection to database.db using sqlite3.connect()
        district: Name of District as String
        institution: Institution data from getInstitutionDetails function
        student_data: Student data from getStudentDetails function

    Returns:
        True: if inserted into DB successfully
        False: if any errors are encountered
    """
    proceed = True

    try:
        cursor = conn.cursor()

        # Insert Institution
        inst_name = institution["name"]
        inst_place = institution["place"]
        inst_number = institution["number"]
        inst_email = institution["email"]

        schoolSQL = """
        INSERT INTO Schools (
            SchoolName,
            District,
            Place,
            Phone,
            Email
        )
        VALUES ( ?, ?, ?, ?, ?)
        """
        values = inst_name, district, inst_place, inst_number, inst_email
        cursor.execute(schoolSQL, values)

        # Get the auto-incremented SchoolID
        school_id = cursor.lastrowid

        # Insert Students
        for key, value in student_data.items():
            name = value[0]
            standard = value[1]
            ifsc = value[2]
            acc_no = value[3]
            holder = value[4]
            branch = value[5]

            studentSQL = """
            INSERT INTO Students (
                SchoolID,
                StudentName,
                Class,
                IFSC,
                AccNo,
                AccHolder,
                Branch
            )
            VALUES ( ?, ?, ?, ?, ?, ?, ?)
            """
            variables = school_id, name, standard, ifsc, acc_no, holder, branch
            cursor.execute(studentSQL, variables)

        print("Commiting Changes")
        conn.commit()

    except IntegrityError as e:
        print(f"IntegrityError: {e}")
        proceed = False
    except Exception as e:
        print(f"Error: {e}")
        proceed = False

    return proceed


# ============================= MAIN FUNCTION ============================== #


if __name__ == "__main__":
    main()
